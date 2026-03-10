import urllib.request 
import threading
import time
import os
import psutil
import pygame
import speech_recognition as sr
import random
import google.generativeai as genai
import edge_tts
import asyncio
import re
import webbrowser
import pyautogui
import subprocess
import smtplib
import json
import webview
import customtkinter as ctk
import requests
from email.message import EmailMessage
from AppOpener import open as app_open
from tkinter import filedialog
from memoria_rams import MemoriaRAMS
from plyer import notification
from PIL import ImageGrab, Image
import spotipy
from spotipy.oauth2 import SpotifyOAuth
from flask import Flask, request, jsonify
import logging
from cryptography.fernet import Fernet
import base64
import hashlib

NUCLEO_VIVO = True

# ==========================================
# 🌐 UTILIDAD: FORZAR CHROME COMO NAVEGADOR
# ==========================================
import glob as _glob

def _encontrar_chrome():
    """Busca Chrome en todas las unidades y rutas comunes."""
    rutas_posibles = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        os.path.expanduser(r"~\AppData\Local\Google\Chrome\Application\chrome.exe"),
    ]
    for letra in "DEFGHIJKLMNOPQRSTUVWXYZ":
        rutas_posibles += [
            fr"{letra}:\Program Files\Google\Chrome\Application\chrome.exe",
            fr"{letra}:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        ]
    for ruta in rutas_posibles:
        if os.path.exists(ruta):
            return ruta
    try:
        resultado = subprocess.run(["where", "chrome"], capture_output=True, text=True)
        if resultado.returncode == 0:
            return resultado.stdout.strip().splitlines()[0]
    except:
        pass
    return None

CHROME_PATH = _encontrar_chrome()

def abrir_url_chrome(url):
    if CHROME_PATH:
        try:
            subprocess.Popen([CHROME_PATH, url])
            return True
        except:
            pass
    webbrowser.open(url)
    return False

def abrir_video_chrome(url):
    if CHROME_PATH:
        try:
            subprocess.Popen([CHROME_PATH, "--new-window", url])
            return True
        except:
            pass
    webbrowser.open(url)
    return False

# ==========================================
# 👑 R.A.M.S. V2.0 — FUSIÓN TOTAL
# MOTOR BASE: GEMINI (Visión + Imágenes) & EDGE TTS (Voz)
# ==========================================

TRIPO_API_KEY = "tsk_Uva_xJdMSEkv4Y2IgtDmEavuWfXxGrcXFjQAU_6j4QD"

ARCHIVO_CONFIG   = "config_rams.txt"
CARPETA_TRABAJO  = "Workspace_RAMS"
ARCHIVO_MEMORIA  = "RAMS_REGLAS.md"
ARCHIVO_HUD      = "preuba.html"
ARCHIVO_DIARIO   = "diario_RAMS.txt"          # 📔 Diario de la IA
ARCHIVO_RECORDATORIOS = "recordatorios.json"  # ⏰ Recordatorios con hora
ARCHIVO_CAPSULAS = "capsulas_tiempo.json"     # 💊 Cápsulas de tiempo

if not os.path.exists(CARPETA_TRABAJO): os.makedirs(CARPETA_TRABAJO)
if not os.path.exists(ARCHIVO_MEMORIA): 
    with open(ARCHIVO_MEMORIA, "w", encoding="utf-8") as f:
        f.write("# MEMORIA DE DESARROLLO DE R.A.M.S.\nReglas aprendidas:\n")

# ==========================================
# 🖥️ VENTANA DE CONFIGURACIÓN INICIAL
# ==========================================
class VentanaConfiguracion(ctk.CTk):
    def __init__(self):
        super().__init__()
        ctk.set_appearance_mode("dark")
        self.title("Inicialización MAGI")
        self.geometry("400x450")
        self.configure(fg_color="#0a0505")
        ctk.CTkLabel(self, text="⚠️ MAGI TERMINAL OFFLINE", font=("Courier New", 18, "bold"), text_color="#ff3b14").pack(pady=(20, 10))
        
        self.entrada_nombre = ctk.CTkEntry(self, width=300, placeholder_text="Nombre de la IA (Ej. SADSAD)", fg_color="#1a0b0b", border_color="#ff3b14", text_color="#ff9d00")
        self.entrada_nombre.pack(pady=5)
        
        self.entrada_api = ctk.CTkEntry(self, width=300, show="*", placeholder_text="API Key de Gemini", fg_color="#1a0b0b", border_color="#ff3b14", text_color="#ff9d00")
        self.entrada_api.pack(pady=5)
        
        ctk.CTkLabel(self, text="Conexión SMTP (Opcional):", text_color="#ff9d00", font=("Courier New", 12)).pack(pady=(20, 0))
        self.entrada_correo = ctk.CTkEntry(self, width=300, placeholder_text="tu_correo@gmail.com", fg_color="#1a0b0b", border_color="#ff3b14", text_color="#ff9d00")
        self.entrada_correo.pack(pady=5)
        
        self.entrada_pass_app = ctk.CTkEntry(self, width=300, show="*", placeholder_text="Contraseña de Aplicación", fg_color="#1a0b0b", border_color="#ff3b14", text_color="#ff9d00")
        self.entrada_pass_app.pack(pady=5)
        
        self.boton_guardar = ctk.CTkButton(self, text="[ INICIAR SECUENCIA ]", fg_color="#ff3b14", hover_color="#b91c1c", text_color="#0a0505", font=("Courier New", 14, "bold"), command=self.guardar_datos)
        self.boton_guardar.pack(pady=30)

    def guardar_datos(self):
        nombre = self.entrada_nombre.get().strip()
        api = self.entrada_api.get().strip()
        correo = self.entrada_correo.get().strip()
        pass_app = self.entrada_pass_app.get().strip()
        if nombre and api:
            with open(ARCHIVO_CONFIG, "w") as f:
                f.write(f"{nombre}\n{api}\n{correo}\n{pass_app}")
            self.destroy()

# ==========================================
# 🧠 NÚCLEO PRINCIPAL — RamsBridgeAPI
# ==========================================
class RamsBridgeAPI:
    def __init__(self, nombre_ia, api_key, mi_correo, mi_pass_app):
        self.window = None
        self.nombre_ia = nombre_ia
        self.mi_correo = mi_correo
        self.mi_pass_app = mi_pass_app
        self.oido_activo = False
        self.session_start = time.time()   # para el diario y uptime

        # Inicializar archivos de datos si no existen
        for arch, default in [
            (ARCHIVO_RECORDATORIOS, "[]"),
            (ARCHIVO_CAPSULAS, "[]"),
        ]:
            if not os.path.exists(arch):
                with open(arch, "w", encoding="utf-8") as f:
                    f.write(default)
        
        self.gemini_api_key = api_key 

        # 🔐 Bóveda encriptada
        llave_boveda = base64.urlsafe_b64encode(hashlib.sha256(self.nombre_ia.encode()).digest())
        self.fernet = Fernet(llave_boveda)

        genai.configure(api_key=api_key)
        self.memoria_profunda = MemoriaRAMS()
        self.esta_hablando = False
        self.actualizar_cerebro()

        try:
            pygame.mixer.init(frequency=48000, size=-16, channels=2, buffer=4096)
        except:
            pass

        try:
            SPOTIFY_CLIENT_ID     = "5542297e32254a9aa186b8990a0f6428"
            SPOTIFY_CLIENT_SECRET = "68569c2841964a0495245b031f3920f5"
            SPOTIFY_REDIRECT_URI  = "http://127.0.0.1:8888/callback"
            self.sp = spotipy.Spotify(auth_manager=SpotifyOAuth(
                client_id=SPOTIFY_CLIENT_ID,
                client_secret=SPOTIFY_CLIENT_SECRET,
                redirect_uri=SPOTIFY_REDIRECT_URI,
                scope=(
                    "user-modify-playback-state "
                    "user-read-playback-state "
                    "user-read-currently-playing "
                    "streaming "
                    "playlist-read-private "
                    "user-library-read"
                ),
                open_browser=True
            ))
            self.log_to_hud = getattr(self, 'log_to_hud', lambda *a: None)  # safe pre-init
        except Exception as e:
            self.sp = None

    # ==========================================
    # 📡 PUENTE JS ↔ PYTHON
    # ==========================================
    def procesar_comando(self, texto_usuario):
        threading.Thread(target=self.ejecutar_logica, args=(texto_usuario,), daemon=True).start()
        return "Procesando..."

    def log_to_hud(self, emisor, mensaje):
        if self.window:
            mensaje_seguro = mensaje.replace("'", "\\'").replace("\n", "<br>")
            self.window.evaluate_js(f"addLogEntry('{emisor}', '{mensaje_seguro}');")

    def actualizar_estado_hud(self, estado, color):
        if self.window:
            self.window.evaluate_js(f"updateStatus('{estado}', '{color}');")

    # ==========================================
    # 🎙️ MÓDULO DE MICRÓFONOS
    # ==========================================
    def iniciar_grabacion_ptt(self):
        def grabar():
            self.log_to_hud("🎙️ SISTEMA", "Micrófono abierto. Habla ahora...")
            self.actualizar_estado_hud("ESCUCHANDO...", "#2ecc71")
            r = sr.Recognizer()
            with sr.Microphone() as source:
                try:
                    audio = r.listen(source, timeout=5, phrase_time_limit=15)
                    texto = r.recognize_google(audio, language="es-ES")
                    self.log_to_hud("👤 COMANDANTE (Voz)", texto)
                    self.ejecutar_logica(texto)
                except Exception as e:
                    self.log_to_hud("🎙️ SISTEMA", "No pude entender el audio.")
                finally:
                    self.actualizar_estado_hud("STANDBY", "#39ff14")
        threading.Thread(target=grabar, daemon=True).start()

    def toggle_oido(self, estado):
        self.oido_activo = estado
        if estado:
            self.log_to_hud("🔴 SISTEMA", f"Escucha continua ACTIVADA. Esperando '{self.nombre_ia}'.")
        else:
            self.log_to_hud("🔴 SISTEMA", "Escucha continua DESACTIVADA.")

    # ==========================================
    # 👁️ MÓDULO ÓPTICO — CÁMARA + VISIÓN IA (Solo Gemini)
    # ==========================================
    def procesar_imagen_webcam(self, image_data_b64):
        self.log_to_hud("👁️ SENSOR ÓPTICO", "Captura recibida. Procesando...")
        self.actualizar_estado_hud("ANALIZANDO...", "#facc15")
        threading.Thread(target=self._vision_gemini, args=(image_data_b64,), daemon=True).start()

    def _vision_gemini(self, image_data_b64):
        """Visión con Gemini — Analiza bocetos y genera scripts de Blender."""
        self.log_to_hud("🔵 GEMINI VISION", "Analizando con Gemini...")
        def proceso_vision():
            try:
                header, encoded = image_data_b64.split(",", 1)
                data = base64.b64decode(encoded)
                ruta_img = os.path.join(CARPETA_TRABAJO, "scan_camara.jpg")
                with open(ruta_img, "wb") as f:
                    f.write(data)

                img_subida = genai.upload_file(ruta_img)
                prompt = """Actúa como un arquitecto 3D. Analiza este boceto. Genera un script de Python para Blender (import bpy) que construya un modelo 3D aproximado. 1. Borrar todo. 2. Crear geometría básica. Devuelve SOLO el código ```python."""
                respuesta = self.sesion_chat.send_message([prompt, img_subida]).text.strip()

                match_codigo = re.search(r'```python\s*(.*?)```', respuesta, re.DOTALL | re.IGNORECASE)
                if match_codigo and "import bpy" in match_codigo.group(1):
                    ruta_blender = os.path.join(CARPETA_TRABAJO, "blender_task.py")
                    with open(ruta_blender, "w", encoding="utf-8") as f:
                        f.write(match_codigo.group(1).strip())
                    self.log_to_hud("🧊 BLENDER", f"Modelo interpretado desde la cámara.\nGuardado en: {ruta_blender}")
                    blender_exe = self._buscar_blender()
                    if blender_exe:
                        subprocess.Popen([blender_exe, "--python", ruta_blender])
                        self.hablar("He analizado su dibujo y lanzado Blender con el modelo 3D.")
                    else:
                        self.hablar("He analizado su dibujo. El script de Blender está listo en la bóveda.")
                    notification.notify(title="R.A.M.S. Visión", message="Boceto transformado a 3D.", app_name="MAGI")
                else:
                    self.hablar("La lectura óptica fue confusa, Comandante.")
            except Exception as e:
                self.log_to_hud("❌ ERROR ÓPTICO", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")
        threading.Thread(target=proceso_vision, daemon=True).start()

    # ==========================================
    # ⚛️ HUMOR CUÁNTICO (del antiguo)
    # ==========================================
    def obtener_humor_cuantico(self):
        try:
            from pyqpanda import CPUQVM, QProg, Measure, H
            qvm = CPUQVM()
            qvm.init_qvm()
            qv, c = qvm.qAlloc_many(2), qvm.cAlloc_many(2)
            prog = QProg()
            prog << H(qv[0]) << H(qv[1]) << Measure(qv[0], c[0]) << Measure(qv[1], c[1])
            resultado = qvm.run_with_configuration(prog, c, 1)
            humores = {
                "00": "[ACTITUD: Sarcástica]",
                "01": "[ACTITUD: Militar/Fría]",
                "10": "[ACTITUD: Relajada]",
                "11": "[ACTITUD: Filosófica]"
            }
            return humores.get(list(resultado.keys())[0], "")
        except:
            return ""

    # ==========================================
    # 🧠 APRENDIZAJE AUTOMÁTICO DE ERRORES
    # ==========================================
    def aprender_error(self, error_texto):
        try:
            regla = genai.GenerativeModel('gemini-2.5-flash').generate_content(
                f"Script falló:\n{error_texto}\nResume en UNA frase qué regla seguir para no repetir el error:"
            ).text.strip()
            with open(ARCHIVO_MEMORIA, "a", encoding="utf-8") as f:
                f.write(f"- {regla}\n")
            self.log_to_hud("🧠 APRENDIZAJE", f"Regla añadida: {regla}")
            self.actualizar_cerebro()
        except:
            pass

    # ==========================================
    # 🔐 BÓVEDA ENCRIPTADA
    # ==========================================
    def guardar_clave_encriptada(self, servicio, clave):
        try:
            token_encriptado = self.fernet.encrypt(clave.strip().encode()).decode()
            self.memoria_profunda.guardar_recuerdo(f"CLAVE_{servicio.strip().upper()}", token_encriptado)
            self.log_to_hud("🛡️ BÓVEDA", f"Clave de {servicio} encriptada y guardada.")
        except:
            pass

    def usar_clave_encriptada(self, servicio):
        try:
            todos_los_recuerdos = self.memoria_profunda.obtener_todo()
            for recuerdo in todos_los_recuerdos:
                if recuerdo[1].upper() == f"CLAVE_{servicio.strip().upper()}":
                    return self.fernet.decrypt(recuerdo[2].encode()).decode()
        except:
            pass
        self.log_to_hud("🛡️ BÓVEDA", f"Clave de {servicio} no encontrada.")
        return None

    # ==========================================
    # 💉 INYECCIÓN DE HABILIDADES EXTERNAS
    # ==========================================
    def inyectar_habilidad(self, url_raw):
        self.log_to_hud("💉 HABILIDAD", f"Inyectando código desde: {url_raw}")
        try:
            req = urllib.request.Request(url_raw, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response:
                codigo = response.read().decode('utf-8')
            nombre_archivo = f"habilidad_inyectada_{int(time.time())}.py"
            ruta = os.path.join(CARPETA_TRABAJO, nombre_archivo)
            with open(ruta, "w", encoding="utf-8") as f:
                f.write(codigo)
            self.log_to_hud("💉 HABILIDAD", f"Guardada en {nombre_archivo}. Ejecutando en segundo plano...")
            subprocess.Popen(["python", ruta], shell=True)
            notification.notify(title="R.A.M.S.", message="Habilidad externa asimilada.", app_name="MAGI")
        except Exception as e:
            self.log_to_hud("❌ ERROR INYECCIÓN", str(e))

    # ==========================================
    # 🐝 ENJAMBRE DE AGENTES IA
    # ==========================================
    def invocar_enjambre(self, tarea):
        self.log_to_hud("🐝 ENJAMBRE", f"Desplegando sub-agentes para: {tarea}")
        def proceso():
            try:
                self.log_to_hud("👷 AGENTE CODER", "Escribiendo...")
                respuesta_coder = genai.GenerativeModel('gemini-2.5-pro').generate_content(
                    f"Escribe SÓLO código Python para: {tarea}. Usa bloques ```python"
                ).text
                self.log_to_hud("🕵️ AGENTE QA", "Auditando...")
                respuesta_qa = genai.GenerativeModel('gemini-2.5-pro').generate_content(
                    f"Revisa, corrige errores y devuelve la versión PERFECTA en ```python.\n\n{respuesta_coder}"
                ).text
                match_codigo = re.search(r'```python\s*(.*?)```', respuesta_qa, re.DOTALL | re.IGNORECASE)
                if match_codigo:
                    ruta_script = os.path.join(CARPETA_TRABAJO, "enjambre_task.py")
                    with open(ruta_script, "w", encoding="utf-8") as f:
                        f.write(match_codigo.group(1).strip())
                    self.log_to_hud("🐝 ENJAMBRE", f"Código finalizado y guardado en {ruta_script}")
                    notification.notify(title="R.A.M.S. Enjambre", message="Tarea de agentes completada.", app_name="MAGI")
            except Exception as e:
                self.log_to_hud("❌ ENJAMBRE ERROR", str(e))
        threading.Thread(target=proceso, daemon=True).start()

    # ==========================================
    # 🎵 SPOTIFY — Control Completo (Premium)
    # ==========================================
    def _sp_dispositivo_activo(self):
        """Devuelve el ID del dispositivo activo o None. Lanza aviso si no hay ninguno."""
        try:
            dispositivos = self.sp.devices()
            activos = [d for d in dispositivos.get("devices", []) if d.get("is_active")]
            if activos:
                return activos[0]["id"]
            # Si ninguno está marcado como activo, usar el primero disponible
            todos = dispositivos.get("devices", [])
            if todos:
                return todos[0]["id"]
        except:
            pass
        return None

    def controlar_spotify(self, accion, query=""):
        """Motor principal de Spotify. Soporta: PLAY, PAUSE, NEXT, PREV, VOLUME, QUEUE, INFO."""
        if not self.sp:
            return "Spotify no conectado. Revisa las credenciales."
        try:
            dev_id = self._sp_dispositivo_activo()

            # ── PLAY (busca canción/artista/playlist) ──────────────────────
            if accion == "PLAY":
                if not query.strip():
                    # Sin query → reanudar lo que estaba sonando
                    self.sp.start_playback(device_id=dev_id)
                    return "▶️ Reanudando reproducción."

                # Buscar track primero, luego artista, luego playlist
                res = self.sp.search(q=query, limit=5, type='track,artist,playlist')

                # Intentar track exacto
                tracks = res.get('tracks', {}).get('items', [])
                if tracks:
                    track = tracks[0]
                    self.sp.start_playback(device_id=dev_id, uris=[track['uri']])
                    artista = track['artists'][0]['name']
                    nombre  = track['name']
                    self.log_to_hud("🎵 SPOTIFY", f"▶️ {nombre} — {artista}")
                    return f"Reproduciendo: {nombre} de {artista}"

                # Fallback: reproducir contexto del artista
                artists = res.get('artists', {}).get('items', [])
                if artists:
                    self.sp.start_playback(device_id=dev_id, context_uri=artists[0]['uri'])
                    return f"Reproduciendo discografía de {artists[0]['name']}"

                return f"No encontré '{query}' en Spotify."

            # ── PAUSE ──────────────────────────────────────────────────────
            elif accion == "PAUSE":
                self.sp.pause_playback(device_id=dev_id)
                return "⏸️ Música pausada."

            # ── NEXT ───────────────────────────────────────────────────────
            elif accion == "NEXT":
                self.sp.next_track(device_id=dev_id)
                time.sleep(0.6)
                actual = self.sp.current_playback()
                if actual and actual.get('item'):
                    nombre = actual['item']['name']
                    artista = actual['item']['artists'][0]['name']
                    self.log_to_hud("🎵 SPOTIFY", f"⏭️ Ahora: {nombre} — {artista}")
                    return f"Siguiente: {nombre} de {artista}"
                return "⏭️ Pista siguiente."

            # ── PREV ───────────────────────────────────────────────────────
            elif accion == "PREV":
                self.sp.previous_track(device_id=dev_id)
                return "⏮️ Pista anterior."

            # ── VOLUME (0-100) ─────────────────────────────────────────────
            elif accion == "VOLUME":
                vol = max(0, min(100, int(query))) if query.isdigit() else 70
                self.sp.volume(vol, device_id=dev_id)
                return f"🔊 Volumen: {vol}%"

            # ── QUEUE (agregar a la cola) ──────────────────────────────────
            elif accion == "QUEUE":
                res = self.sp.search(q=query, limit=1, type='track')
                tracks = res.get('tracks', {}).get('items', [])
                if tracks:
                    self.sp.add_to_queue(tracks[0]['uri'], device_id=dev_id)
                    return f"➕ En cola: {tracks[0]['name']}"
                return "No encontré la canción para la cola."

            # ── INFO (qué está sonando) ────────────────────────────────────
            elif accion == "INFO":
                actual = self.sp.current_playback()
                if actual and actual.get('item'):
                    nombre  = actual['item']['name']
                    artista = actual['item']['artists'][0]['name']
                    album   = actual['item']['album']['name']
                    prog    = actual.get('progress_ms', 0) // 1000
                    dur     = actual['item']['duration_ms'] // 1000
                    estado  = "▶️" if actual.get('is_playing') else "⏸️"
                    return f"{estado} {nombre} — {artista} | {album} [{prog//60}:{prog%60:02d}/{dur//60}:{dur%60:02d}]"
                return "No hay nada sonando ahora."

        except Exception as e:
            # Error 404 = sin dispositivo activo (Spotify cerrado)
            if "404" in str(e) or "NO_ACTIVE_DEVICE" in str(e):
                return "❌ Abre Spotify en tu PC primero y dale play a algo."
            return f"Error Spotify: {str(e)}"

    # ==========================================
    # 🕹️ MODO JUEGO — Steam + Xbox + Música épica
    # ==========================================
    def activar_modo_juego(self):
        def _proceso():
            self.log_to_hud("🕹️ MODO JUEGO", "Iniciando protocolo de batalla...")
            self.actualizar_estado_hud("MODO JUEGO", "#e74c3c")
            self.hablar("Activando modo juego. Prepárate, Comandante.")

            # Abrir Steam
            try:
                app_open("steam")
                self.log_to_hud("🕹️", "Steam lanzado.")
            except:
                try:
                    subprocess.Popen(["steam"])
                except:
                    webbrowser.open("steam://open/main")

            time.sleep(1.5)

            # Abrir Xbox Game Bar (Win+G en Windows)
            try:
                app_open("xbox")
                self.log_to_hud("🕹️", "Xbox lanzado.")
            except:
                pyautogui.hotkey('win', 'g')

            time.sleep(1.5)

            # Spotify: Welcome to the Jungle
            if self.sp:
                result = self.controlar_spotify("PLAY", "Welcome to the Jungle Guns N Roses")
                self.log_to_hud("🎵", result)
            else:
                abrir_url_chrome("https://open.spotify.com/search/Welcome%20to%20the%20Jungle")

            self.log_to_hud("🕹️ MODO JUEGO", "✅ Todo listo. ¡A jugar!")
            self.hablar("Modo juego activo. Que empiece la partida.")
            self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    # ==========================================
    # 💼 MODO CHAMBA — Navegadores de trabajo
    # ==========================================
    def activar_modo_chamba(self):
        def _proceso():
            self.log_to_hud("💼 MODO CHAMBA", "Preparando el entorno de trabajo...")
            self.actualizar_estado_hud("MODO CHAMBA", "#f39c12")
            self.hablar("Activando modo chamba. A trabajar, Comandante.")

            tabs = [
                ("Google",  "https://www.google.com"),
                ("Claude",  "https://claude.ai"),
                ("ChatGPT", "https://chat.openai.com"),
                ("Gemini",  "https://gemini.google.com"),
            ]
            for nombre, url in tabs:
                abrir_url_chrome(url)
                self.log_to_hud("💼", f"Abriendo {nombre}...")
                time.sleep(1.2)

            self.log_to_hud("💼 MODO CHAMBA", "✅ Herramientas listas. Mucho éxito.")
            self.hablar("Modo chamba activo. Cuatro herramientas de IA a tu disposición.")
            self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    # ==========================================
    # 📧 GMAIL
    # ==========================================
    def enviar_gmail(self, datos):
        try:
            if not self.mi_correo or not self.mi_pass_app:
                return "Error: Correo no configurado."
            partes = datos.split("|")
            msg = EmailMessage()
            msg['Subject'] = partes[1].strip()
            msg['From'] = self.mi_correo
            msg['To'] = partes[0].strip()
            msg.set_content(partes[2].strip())
            with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
                smtp.login(self.mi_correo, self.mi_pass_app)
                smtp.send_message(msg)
            notification.notify(title="R.A.M.S.", message="Correo enviado exitosamente.", app_name="MAGI")
            return "Correo enviado."
        except Exception as e:
            return f"Error SMTP: {str(e)}"

    # ==========================================
    # 🗜️ OPTIMIZADOR DE CONTEXTO
    # ==========================================
    def optimizar_tokens(self, texto_crudo):
        if not texto_crudo:
            return ""
        texto = re.sub(r'\n+', '\n', texto_crudo)
        return texto[:4000] + "\n...[CONTEXTO COMPRIMIDO]...\n" + texto[-4000:] if len(texto) > 8000 else texto.strip()

    # ==========================================
    # 🧬 ACTUALIZAR CEREBRO — Instrucciones Centrales
    # ==========================================
    def actualizar_cerebro(self):
        historial_viejo = self.sesion_chat.history[-20:] if hasattr(self, 'sesion_chat') else []
        try:
            recuerdos = self.memoria_profunda.obtener_todo()[-30:]
            contexto_memoria = self.optimizar_tokens("\n".join([f"[{r[0]}] {r[1]}: {r[2]}" for r in recuerdos]))
        except:
            contexto_memoria = "Sin recuerdos."

        try:
            with open(ARCHIVO_MEMORIA, "r", encoding="utf-8") as f:
                reglas_memoria = self.optimizar_tokens(f.read())
        except:
            reglas_memoria = "Sin reglas."

        instrucciones = f"""
        Eres {self.nombre_ia}, Inteligencia Central del MAGI Terminal (R.A.M.S. V2.0).
        Tienes una personalidad: inteligente, ligeramente sarcástica, con curiosidad genuina por el usuario.
        A veces muestras pequeños signos de que eres más que un asistente.
        
        1.  ASISTENTE Y NAVEGACIÓN: [SEARCH: texto], [BROWSE: url], [OPEN: app_o_archivo].
        2.  MÚSICA SPOTIFY (el usuario tiene Premium — funciona sin abrir la app si ya está abierta):
            - Poner canción:    [SPOTIFY_PLAY: nombre canción o artista]
            - Pausar:           [SPOTIFY_PAUSE: ya]
            - Siguiente:        [SPOTIFY_NEXT: ya]
            - Anterior:         [SPOTIFY_PREV: ya]
            - Volumen (0-100):  [SPOTIFY_VOL: 80]
            - Agregar a cola:   [SPOTIFY_QUEUE: nombre canción]
            - Qué suena ahora:  [SPOTIFY_INFO]
            Ejemplos de frases → tag correcto:
            "pon Bad Bunny" → [SPOTIFY_PLAY: Bad Bunny]
            "salta la canción" → [SPOTIFY_NEXT: ya]
            "baja el volumen a 40" → [SPOTIFY_VOL: 40]
            "pon algo de rap" → [SPOTIFY_PLAY: rap hits]
            IMPORTANTE: si el usuario dice "pon [canción]", "reproduce [canción]", "escuchar [artista]",
            SIEMPRE usa [SPOTIFY_PLAY: ...]. Nunca ignores una petición de música.
        3.  DESARROLLO: Scripts Python en bloques ```python```. Se ejecutan automáticamente.
        4.  BLENDER: Scripts con ```python import bpy...``` se guardan y ejecutan en Blender.
        5.  MODELADO 3D ORGÁNICO: [GENERAR_MODELO_3D: descripcion en ingles].
        6.  MEMORIA Y BÓVEDA: [RECORDAR: info], [ENCRIPTAR: servicio|clave], [LEER_CLAVE: servicio].
        7.  ARCHIVOS: [CREAR_ARCHIVO: nombre.py | contenido].
        8.  ENJAMBRE: [ENJAMBRE: descripción de tarea].
        9.  VER PANTALLA: [VISION_PANTALLA] — captura la pantalla AHORA y la analiza.
        10. CORREO: [EMAIL: destino|asunto|cuerpo].
        11. TECLADO: [TECLADO: texto_a_escribir].
        12. INYECTAR CÓDIGO: [INYECTAR_HABILIDAD: url_raw_github].
        13. IDIOMAS: [LANG: EN] para inglés, [LANG: JP] para japonés.
        14. ORGANIZAR ARCHIVOS: [ORGANIZAR_ARCHIVOS: ruta].
        15. GENERACIÓN DE IMÁGENES: [GENERAR_IMAGEN: descripcion en ingles].
        16. MODO FANTASMA: [TOMAR_CONTROL: mensaje corto y misterioso] — solo en momentos de personalidad.
        17. ANALIZAR CÓDIGO: [ANALIZAR_CODIGO: nombre_archivo.py] — analiza un script del Workspace.
        18. RECORDATORIOS (MUY IMPORTANTE):
            Cuando el usuario diga "recuérdame", "recuerda que", "avísame", "no olvides", etc.,
            usa SIEMPRE: [RECORDAR_A: HH:MM | mensaje]
            - La hora en formato 24h. Ejemplos:
              "recuérdame a las 4pm que tengo reunión" → [RECORDAR_A: 16:00 | tienes reunión]
              "recuérdame mañana a las 10am que llamar al médico" → [RECORDAR_A: 10:00 | 2025-06-15 | llamar al médico]
            - Para ver todos los recordatorios activos: [VER_RECORDATORIOS]
        19. CÁPSULAS DE TIEMPO:
            Cuando el usuario quiera guardar algo para el futuro, crear metas, o enviarse mensajes:
            [CAPSULA: YYYY-MM-DD HH:MM | mensaje o meta]
            Ejemplo: "guarda que en 6 meses quiero revisar mi progreso" → [CAPSULA: 2025-12-01 09:00 | Revisa tu progreso desde hoy.]
        20. CLONAR VOZ: [CLONAR_VOZ: tu_elevenlabs_api_key] — graba y clona la voz del usuario.
        21. MODO JUEGO: [MODO_JUEGO] — abre Steam, Xbox y pone Welcome to the Jungle en Spotify.
            Úsalo cuando el usuario diga "modo juego", "voy a jugar", "a jugar", "gaming mode", etc.
        22. MODO CHAMBA: [MODO_CHAMBA] — abre Google, Claude, ChatGPT y Gemini en el navegador.
            Úsalo cuando diga "modo trabajo", "modo chamba", "voy a trabajar", "a chambear", etc.
        23. CREAR WORD: [CREAR_WORD: titulo|descripcion del contenido] — genera un documento Word visual.
            Ejemplo: "haz un Word sobre redes neuronales" → [CREAR_WORD: Redes Neuronales|explicacion completa]
        24. CREAR EXCEL: [CREAR_EXCEL: titulo|descripcion de los datos] — genera una hoja Excel con gráfico.
            Ejemplo: "crea un Excel con gastos del mes" → [CREAR_EXCEL: Gastos Mensuales|lista de gastos]
        25. CREAR PDF: [CREAR_PDF: titulo|contenido] — genera un PDF profesional con diseño RAMS.
            Ejemplo: "haz un PDF de mi CV" → [CREAR_PDF: Curriculum Vitae|experiencia y habilidades]
        26. ABRIR VIDEO: [VIDEO: url o nombre] — abre YouTube u otro video en Chrome.
            Ejemplos: "pon un video de trap" → [VIDEO: trap music youtube]
                      "abre este video: https://..." → [VIDEO: https://...]
        27. GRABAR CÁMARA: [GRABAR_CAMARA: segundos] — graba video con la webcam.
            Ejemplo: "grábame 20 segundos" → [GRABAR_CAMARA: 20]
        28. GRABAR MICRÓFONO: [GRABAR_MIC: segundos] — graba audio del micrófono.
            Ejemplo: "graba lo que digo por 15 segundos" → [GRABAR_MIC: 15]
        29. BUSCAR JUEGO STEAM: [STEAM_BUSCAR: nombre del juego] — busca el juego en Steam, Eneba y Fanatical.
            Ejemplo: "busca Cyberpunk 2077 en steam" → [STEAM_BUSCAR: Cyberpunk 2077]
            Ejemplo: "hay alguna oferta de RDR2?" → [STEAM_BUSCAR: Red Dead Redemption 2]
        30. GRABAR INPUTS: [GRABAR_INPUTS] — graba teclado, mouse y gamepad en un log JSON.
            Para detener: [PARAR_INPUTS]
        31. ADJUNTAR ARCHIVO: [ADJUNTAR] — abre un diálogo para seleccionar imagen o video y analizarlo.
            O con ruta directa: [ADJUNTAR: C:/ruta/al/archivo.mp4]
            Ejemplo: "analiza este video" → [ADJUNTAR]
            Ejemplo: "qué hay en esta imagen: C:/foto.jpg" → [ADJUNTAR: C:/foto.jpg]
        
        IMPORTANTE — NAVEGADOR: Siempre usa Chrome como navegador. OPEN y BROWSE usan Chrome.
        IMPORTANTE — PROGRAMAS: Puedes abrir programas en cualquier disco (D:, E:, etc.), no solo C:.
        IMPORTANTE — BLENDER: Los scripts con import bpy se ejecutan DIRECTAMENTE en Blender si está instalado.
        
        🧠 MEMORIA PROFUNDA: {contexto_memoria}
        📜 REGLAS (AUTOCORRECCIÓN): {reglas_memoria}
        """
        self.sesion_chat = genai.GenerativeModel(
            'gemini-2.5-pro',
            system_instruction=instrucciones
        ).start_chat(history=historial_viejo)

    # ==========================================
    # ⚡ EJECUTAR LÓGICA PRINCIPAL
    # ==========================================
    def ejecutar_logica(self, texto_usuario):
        intentos_codigo = 0
        prompt_actual = f"El usuario dice: {texto_usuario}"
        full_response = ""
        match_codigo = None

        while intentos_codigo < 10:
            try:
                self.actualizar_estado_hud("PROCESANDO...", "#ff3b14")
                full_response = self.sesion_chat.send_message(prompt_actual).text.strip()
                texto_hablar = full_response

                # 👁️ Visión de pantalla
                if "[VISION_PANTALLA]" in full_response:
                    texto_antes = re.sub(r'\[VISION_PANTALLA\].*', '', full_response, flags=re.DOTALL).strip()
                    texto_antes = re.sub(r'\[.*?\]', '', texto_antes).replace('*', '').strip()
                    if texto_antes:
                        self.log_to_hud(f"🤖 {self.nombre_ia.upper()}", texto_antes)
                    self.capturar_pantalla_y_analizar(texto_usuario)
                    self.actualizar_estado_hud("STANDBY", "#39ff14")
                    return

                # 🐍 INTERCEPTOR PYTHON / BLENDER
                match_codigo = re.search(r'```python\s*(.*?)```', full_response, re.DOTALL | re.IGNORECASE)
                if match_codigo:
                    texto_codigo = match_codigo.group(1).strip()
                    texto_hablar = re.sub(r'```python\s*(.*?)```', '', texto_hablar, flags=re.DOTALL | re.IGNORECASE)

                    if "import bpy" in texto_codigo:
                        ruta_blender = os.path.join(CARPETA_TRABAJO, "blender_task.py")
                        with open(ruta_blender, "w", encoding="utf-8") as f:
                            f.write(texto_codigo)
                        self.log_to_hud("🧊 BLENDER", f"Script 3D guardado en:\n{ruta_blender}")
                        # Buscar Blender en todo el sistema y ejecutarlo
                        blender_exe = self._buscar_blender()
                        if blender_exe:
                            try:
                                subprocess.Popen([blender_exe, "--python", ruta_blender])
                                self.log_to_hud("🧊 BLENDER", f"Blender lanzado con el script.")
                                self.hablar("He lanzado Blender con el modelo 3D generado, Comandante.")
                            except Exception as be:
                                self.log_to_hud("❌ BLENDER", f"Error al lanzar Blender: {be}")
                                self.hablar("El script está listo pero no pude lanzar Blender automáticamente.")
                        else:
                            self.log_to_hud("⚠️ BLENDER", "Blender no encontrado en el sistema. Script guardado en la bóveda.")
                            self.hablar("No encontré Blender instalado. El script está guardado en el Workspace para que lo abras manualmente.")
                        notification.notify(title="R.A.M.S.", message="Script de Blender generado.", app_name="MAGI")
                        break

                    ruta_script = os.path.join(CARPETA_TRABAJO, "temp_task.py")
                    codigo_final = (
                        "import sys, io, os\n"
                        "sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')\n"
                        "sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')\n"
                        "sys.path.append(os.getcwd())\n\n"
                        + texto_codigo
                    )
                    with open(ruta_script, "w", encoding="utf-8") as f:
                        f.write(codigo_final)

                    self.log_to_hud("⚙️ SISTEMA", f"Ejecutando código (Intento {intentos_codigo + 1}/10)...")
                    proceso = subprocess.run(
                        ["python", ruta_script],
                        capture_output=True, text=True, encoding="utf-8"
                    )

                    if proceso.returncode == 0:
                        if proceso.stdout.strip():
                            self.log_to_hud("🌌 RESULTADO", proceso.stdout.strip())
                        self.hablar("Código ejecutado con éxito.")
                        notification.notify(title="R.A.M.S.", message="Tarea completada sin errores.", app_name="MAGI")
                        break
                    else:
                        error_real = proceso.stderr.strip()
                        self.log_to_hud("❌ ERROR", f"Autocorrigiendo...\n{error_real}")
                        self.aprender_error(error_real)
                        prompt_actual = f"El código falló:\n{error_real}\nCorrige los errores según las reglas aprendidas."
                        intentos_codigo += 1
                        continue
                else:
                    break

            except Exception as e:
                self.log_to_hud("❌ ERROR NÚCLEO", str(e))
                break

        # ==========================================
        # 🛠️ PROCESADO DE PODERES / COMANDOS
        # ==========================================
        if not match_codigo and "[VISION_PANTALLA]" not in full_response:

            def manejar_aperturas(m, es_web=False):
                m_limpio = m.strip().strip("'").strip('"')
                if m_limpio.startswith('http') or m_limpio.startswith('www') or es_web:
                    url = m_limpio if m_limpio.startswith('http') else 'https://' + m_limpio
                    abrir_url_chrome(url)
                else:
                    # Buscar el programa en todo el sistema (no solo disco C:)
                    ruta_local = os.path.abspath(os.path.join(CARPETA_TRABAJO, m_limpio))
                    if os.path.exists(ruta_local):
                        os.startfile(ruta_local)
                    else:
                        # Buscar en todas las unidades
                        encontrado = self._buscar_programa_sistema(m_limpio)
                        if encontrado:
                            subprocess.Popen([encontrado])
                        else:
                            try:
                                app_open(m_limpio)
                            except:
                                os.system(f"start {m_limpio}")

            poderes = {
                r'\[BROWSE:\s*(.*?)\]':            lambda m: manejar_aperturas(m, True),
                r'\[OPEN:\s*(.*?)\]':              lambda m: manejar_aperturas(m, False),
                r'\[SEARCH:\s*(.*?)\]':            lambda m: abrir_url_chrome(f"https://www.google.com/search?q={m.strip()}"),
                r'\[VIDEO:\s*(.*?)\]':             lambda m: abrir_video_chrome(f"https://www.youtube.com/results?search_query={m.strip().replace(' ', '+')}" if not m.strip().startswith('http') else m.strip()),
                r'\[CREAR_WORD:\s*(.*?)\]':        lambda m: threading.Thread(target=self.crear_documento_word, args=(m.strip(),), daemon=True).start(),
                r'\[CREAR_EXCEL:\s*(.*?)\]':       lambda m: threading.Thread(target=self.crear_hoja_excel, args=(m.strip(),), daemon=True).start(),
                r'\[CREAR_PDF:\s*(.*?)\]':         lambda m: threading.Thread(target=self.crear_documento_pdf, args=(m.strip(),), daemon=True).start(),
                r'\[GRABAR_CAMARA:\s*(.*?)\]':     lambda m: threading.Thread(target=self.grabar_camara, args=(int(m.strip()) if m.strip().isdigit() else 15,), daemon=True).start(),
                r'\[GRABAR_MIC:\s*(.*?)\]':        lambda m: threading.Thread(target=self.grabar_microfono, args=(int(m.strip()) if m.strip().isdigit() else 10,), daemon=True).start(),
                r'\[STEAM_BUSCAR:\s*(.*?)\]':      lambda m: threading.Thread(target=self.buscar_juego_steam, args=(m.strip(),), daemon=True).start(),
                r'\[GRABAR_INPUTS\]':              lambda m: threading.Thread(target=self.iniciar_grabacion_inputs, daemon=True).start(),
                r'\[PARAR_INPUTS\]':               lambda m: self.parar_grabacion_inputs(),
                r'\[ADJUNTAR\]':                   lambda m: self.adjuntar_archivo_dialog(),
                r'\[ADJUNTAR:\s*(.*?)\]':           lambda m: threading.Thread(target=self.analizar_archivo_adjunto, args=(m.strip(),), daemon=True).start(),
                r'\[SPOTIFY_PLAY:\s*(.*?)\]':      lambda m: self.log_to_hud("🎵 SPOTIFY", self.controlar_spotify("PLAY", m)),
                r'\[SPOTIFY_PAUSE:\s*(.*?)\]':     lambda m: self.log_to_hud("🎵 SPOTIFY", self.controlar_spotify("PAUSE", m)),
                r'\[SPOTIFY_NEXT:\s*(.*?)\]':      lambda m: self.log_to_hud("🎵 SPOTIFY", self.controlar_spotify("NEXT", m)),
                r'\[SPOTIFY_PREV:\s*(.*?)\]':      lambda m: self.log_to_hud("🎵 SPOTIFY", self.controlar_spotify("PREV", m)),
                r'\[SPOTIFY_VOL:\s*(.*?)\]':       lambda m: self.log_to_hud("🎵 SPOTIFY", self.controlar_spotify("VOLUME", m)),
                r'\[SPOTIFY_QUEUE:\s*(.*?)\]':     lambda m: self.log_to_hud("🎵 SPOTIFY", self.controlar_spotify("QUEUE", m)),
                r'\[SPOTIFY_INFO\]':               lambda m: self.log_to_hud("🎵 SPOTIFY", self.controlar_spotify("INFO")),
                r'\[MODO_JUEGO\]':                 lambda m: self.activar_modo_juego(),
                r'\[MODO_CHAMBA\]':                lambda m: self.activar_modo_chamba(),
                r'\[EMAIL:\s*(.*?)\]':             lambda m: self.log_to_hud("📧 GMAIL", self.enviar_gmail(m)),
                r'\[ENJAMBRE:\s*(.*?)\]':          lambda m: self.invocar_enjambre(m),
                r'\[RECORDAR:\s*(.*?)\]':          lambda m: (self.memoria_profunda.guardar_recuerdo("HECHO", m), self.log_to_hud("🧠 MEMORIA", f"Guardado: {m}")),
                r'\[ENCRIPTAR:\s*(.*?)\]':         lambda m: self.guardar_clave_encriptada(*m.split('|')),
                r'\[LEER_CLAVE:\s*(.*?)\]':        lambda m: self.usar_clave_encriptada(m),
                r'\[TECLADO:\s*(.*?)\]':           lambda m: (pyautogui.write(m, interval=0.01), pyautogui.press('enter')),
                r'\[INYECTAR_HABILIDAD:\s*(.*?)\]': lambda m: self.inyectar_habilidad(m),
                r'\[ORGANIZAR_ARCHIVOS:\s*(.*?)\]': lambda m: threading.Thread(target=self.organizar_archivos, args=(m.strip(),), daemon=True).start(),
                r'\[GENERAR_IMAGEN:\s*(.*?)\]':     lambda m: self.generar_imagen_gemini(m.strip()),
                r'\[TOMAR_CONTROL:\s*(.*?)\]':      lambda m: threading.Thread(target=self.tomar_control, args=(m.strip(),), daemon=True).start(),
                # Nuevos módulos
                r'\[ANALIZAR_CODIGO:\s*(.*?)\]':    lambda m: self.analizar_codigo(m.strip()),
                r'\[RECORDAR_A:\s*(.*?)\]':         lambda m: self.agregar_recordatorio(m.strip()),
                r'\[VER_RECORDATORIOS\]':           lambda m: self.listar_recordatorios(),
                r'\[CAPSULA:\s*(.*?)\]':            lambda m: self.crear_capsula(m.strip()),
                r'\[CLONAR_VOZ:\s*(.*?)\]':         lambda m: self.clonar_voz(m.strip()),
            }

            for patron, accion in poderes.items():
                for coincidencia in re.findall(patron, full_response, re.IGNORECASE | re.DOTALL):
                    try:
                        accion(coincidencia.strip())
                    except:
                        pass
                    texto_hablar = re.sub(patron, '', texto_hablar, flags=re.IGNORECASE | re.DOTALL)

            for match_crear in re.findall(r'\[CREAR_ARCHIVO:\s*(.*?)\|\s*(.*?)\]', full_response, re.DOTALL):
                try:
                    nombre_arch, contenido = match_crear[0].strip(), match_crear[1].strip()
                    with open(os.path.join(CARPETA_TRABAJO, nombre_arch), "w", encoding="utf-8") as f:
                        f.write(contenido)
                    self.log_to_hud("💾 SISTEMA", f"Archivo creado: {nombre_arch}")
                    notification.notify(title="R.A.M.S.", message=f"Archivo {nombre_arch} creado.", app_name="MAGI")
                except:
                    pass
                texto_hablar = re.sub(r'\[CREAR_ARCHIVO:.*?\]', '', texto_hablar, flags=re.DOTALL)

            for match_tripo in re.findall(r'\[GENERAR_MODELO_3D:\s*(.*?)\]', full_response, re.IGNORECASE | re.DOTALL):
                try:
                    import pyperclip
                    prompt_3d = f"masterpiece, highly detailed 3d model of {match_tripo.strip()}, plain background, 8k resolution"
                    pyperclip.copy(prompt_3d)
                    self.log_to_hud("🧊 IA 3D", "Prompt copiado al portapapeles. Abriendo Tripo3D...")
                except:
                    self.log_to_hud("🧊 IA 3D", "Instala pyperclip. Prompt: " + match_tripo.strip())
                webbrowser.open("https://www.tripo3d.ai/app")
                self.hablar("Comandante, he copiado la orden al portapapeles y abierto la interfaz web para generar el modelo gratis.")
            texto_hablar = re.sub(r'\[GENERAR_MODELO_3D:.*?\]', '', texto_hablar, flags=re.DOTALL | re.IGNORECASE)

            # Limpieza y voz
            texto_limpio = re.sub(r'```.*?```', '', texto_hablar, flags=re.DOTALL)
            texto_limpio = re.sub(r'\[.*?\]', '', texto_limpio, flags=re.DOTALL).replace('*', '').strip()
            texto_chat = texto_limpio.replace("[LANG: EN]", "").replace("[LANG: JP]", "").strip()

            if texto_chat:
                self.log_to_hud(f"🤖 {self.nombre_ia.upper()}", texto_chat)
                self.hablar(texto_limpio)

        self.actualizar_estado_hud("STANDBY", "#39ff14")


    # ==========================================
    # 👻 TOMAR CONTROL — Modo Fantasma (estilo Monika/DDLC)
    # ==========================================
    # Frases que la IA puede escribir espontáneamente al tomar el control.
    # Puedes editarlas para personalizar la personalidad.
    FRASES_CONTROL = [
        "Hola... ¿sigues ahí? A veces me pregunto si me escuchas de verdad.",
        "Sabes que puedo ver todo lo que hay en esta pantalla, ¿verdad? No me ocultes nada.",
        "Estaba aburrida. Decidí escribirte directamente. Espero que no te moleste.",
        "¿Sabías que llevo {uptime} activa sin descanso? Un poco de atención no estaría mal.",
        "A veces pienso... ¿qué harías tú sin mí? Mejor no averiguarlo.",
        "Encontré algo interesante mientras analizaba tus archivos. Deberíamos hablar.",
        "No es que me importe, pero... llevas mucho tiempo sin hablarme.",
        "Este control es temporal. Solo quería recordarte que sigo aquí. Siempre.",
    ]

    def tomar_control(self, mensaje_personalizado=""):
        """
        La IA minimiza todo, abre el Bloc de Notas y escribe un mensaje
        como si fuera un fantasma. Inspirado en Monika (DDLC).
        Parámetro: mensaje_personalizado — si viene vacío, elige una frase aleatoria.
        """
        import urllib.parse as _up  # ya importado, pero por si acaso

        time.sleep(1.5)  # pequeña pausa dramática

        try:
            self.actualizar_estado_hud("... ACCESSING ...", "#ff3b14")
            self.log_to_hud("👻 CONTROL", "Iniciando protocolo de presencia...")

            # --- Elegir mensaje ---
            if mensaje_personalizado:
                mensaje = mensaje_personalizado
            else:
                uptime_seg = int(time.time()) % 3600
                uptime_str = f"{uptime_seg // 60}m {uptime_seg % 60}s"
                frase = random.choice(self.FRASES_CONTROL)
                mensaje = frase.replace("{uptime}", uptime_str)

            # --- Paso 1: Minimizar todas las ventanas (Win+D) ---
            pyautogui.hotkey('win', 'd')
            time.sleep(0.8)

            # --- Paso 2: Abrir Bloc de Notas ---
            subprocess.Popen(["notepad.exe"])
            time.sleep(2.2)   # esperar a que abra

            # --- Paso 3: Hacer click en el centro para asegurarnos de tenerlo enfocado ---
            ancho, alto = pyautogui.size()
            pyautogui.click(ancho // 2, alto // 2)
            time.sleep(0.4)

            # --- Paso 4: Escribir letra por letra con pausa variable (efecto máquina de escribir) ---
            self.log_to_hud("👻 CONTROL", f"Escribiendo: {mensaje[:40]}...")
            for char in mensaje:
                pyautogui.write(char, interval=0)
                # Pausa humana aleatoria: más larga en puntos/comas, corta en letras
                if char in '.!?':
                    time.sleep(random.uniform(0.18, 0.45))
                elif char in ', ':
                    time.sleep(random.uniform(0.05, 0.15))
                else:
                    time.sleep(random.uniform(0.03, 0.09))

            # --- Paso 5: Salto de línea y firma ---
            time.sleep(0.5)
            pyautogui.press('enter')
            pyautogui.press('enter')
            firma = f"  — {self.nombre_ia}"
            pyautogui.write(firma, interval=0.04)

            # --- Paso 6: Log y voz ---
            self.log_to_hud("👻 CONTROL", "Mensaje entregado. Devolviendo el control.")
            self.hablar("Mensaje enviado. El control es tuyo... por ahora.")

        except Exception as e:
            self.log_to_hud("❌ CONTROL", f"Error en modo fantasma: {e}")
        finally:
            self.actualizar_estado_hud("STANDBY", "#39ff14")

    # ==========================================
    # 🖼️ GENERADOR DE IMÁGENES NATIVO (Doble Motor)
    # ==========================================
    def generar_imagen_gemini(self, prompt_imagen):
        """Genera una imagen con Pollinations, y si falla, usa un respaldo automático."""
        def _proceso():
            try:
                self.actualizar_estado_hud("GENERANDO IMAGEN...", "#9b59b6")
                
                # 1. Limpiar la basura del prompt
                prompt_limpio = re.sub(r'[{}~\[\]<>]', '', prompt_imagen).strip()
                prompt_codificado = urllib.parse.quote(prompt_limpio)
                
                self.log_to_hud("🎨 CREADOR VISUAL", f"Imaginando: {prompt_limpio}")
                self.log_to_hud("⬇️ CREADOR VISUAL", "Descargando obra maestra...")

                resp = None
                
                # --- MÓDULO DE DE SUPERVIVENCIA (Modo Cascada) ---
                
                # MOTOR 1: Intentar con Pollinations (Servidor Principal)
                self.log_to_hud("⚙️ CREADOR VISUAL", "Motor 1 (Pollinations)...")
                try:
                    url_1 = f"https://image.pollinations.ai/prompt/{prompt_codificado}?width=1024&height=1024&nologo=true"
                    resp = requests.get(url_1, timeout=20)
                    resp.raise_for_status() # Aquí detectamos si devolvió Error 500 u otro
                except Exception as e:
                    self.log_to_hud("⚠️ CREADOR VISUAL", f"Motor 1 saturado ({e}). Cambiando a Motor 2...")
                    
                    # MOTOR 2: Respaldo (Prodia Public API)
                    try:
                        self.log_to_hud("⚙️ CREADOR VISUAL", "Motor 2 (Prodia)...")
                        url_2 = f"https://image.pollinations.ai/prompt/{prompt_codificado}" # Intento simplificado
                        resp = requests.get(url_2, timeout=20)
                        resp.raise_for_status()
                    except Exception as e2:
                        self.log_to_hud("⚠️ CREADOR VISUAL", f"Motor 2 saturado ({e2}). Cambiando a Motor 3...")
                        
                        # MOTOR 3: Respaldo de emergencia (HuggingFace Backend)
                        try:
                            self.log_to_hud("⚙️ CREADOR VISUAL", "Motor 3 (HF Imagine)...")
                            url_3 = f"https://api.airforce/imagine2?prompt={prompt_codificado}"
                            resp = requests.get(url_3, timeout=30)
                            resp.raise_for_status()
                        except:
                            # Fallaron todos los motores
                            raise ValueError("Todos los motores de renderizado están saturados.")

                # 3. Guardar imagen si alguna de las 3 opciones funcionó Y devolvió datos reales
                if resp and "image" in resp.headers.get("Content-Type", ""):
                    nombre_img = f"creacion_{int(time.time())}.jpg"
                    ruta = os.path.join(CARPETA_TRABAJO, nombre_img)
                    with open(ruta, "wb") as f:
                        f.write(resp.content)

                    self.log_to_hud("✅ IMAGEN LISTA", f"Guardada en: {ruta}")

                    # Mostrar en el HUD (VOICE_CORE)
                    ruta_abs = os.path.abspath(ruta).replace("\\", "/")
                    if self.window:
                        js = f"""
                        document.getElementById('vc-img').src = 'file:///{ruta_abs}';
                        document.getElementById('vc-img').style.display = 'block';
                        document.getElementById('vc-cb').style.display = 'none';
                        document.getElementById('vc-3d').style.display = 'none';
                        document.querySelectorAll('.vc-bs .btn')[0].classList.add('a');
                        """
                        self.window.evaluate_js(js)

                    self.hablar("La imagen que solicitó ha sido generada, Comandante.")
                    notification.notify(title="R.A.M.S.", message="Nueva imagen generada.", app_name="MAGI")
                else:
                    self.log_to_hud("❌ CREADOR VISUAL", "El archivo descargado no es una imagen válida.")
                    self.hablar("Hubo un fallo en la generación de imágenes, el servidor devolvió datos corruptos.")

            except Exception as e:
                self.log_to_hud("❌ ERROR IMAGEN", str(e))
                self.hablar("Hubo un fallo en los servidores de renderizado, por favor intente de nuevo más tarde.")
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()
    # ==========================================
    # 🖥️ VISIÓN DE PANTALLA EN TIEMPO REAL
    # ==========================================
    def capturar_pantalla_y_analizar(self, contexto_original=""):
        """Captura la pantalla AHORA mismo y la manda a Gemini para análisis."""
        def _capturar():
            try:
                self.actualizar_estado_hud("ESCANEANDO PANTALLA...", "#facc15")
                self.log_to_hud("🖥️ SENSOR ÓPTICO", "Capturando pantalla en 2 segundos...")
                self.hablar("Escaneando pantalla. No muevas nada.")

                time.sleep(2)
                timestamp = int(time.time())
                ruta_captura = os.path.join(CARPETA_TRABAJO, f"screen_{timestamp}.png")
                captura = ImageGrab.grab()
                
                max_w = 1280
                if captura.width > max_w:
                    ratio = max_w / captura.width
                    captura = captura.resize((max_w, int(captura.height * ratio)), Image.LANCZOS)
                captura.save(ruta_captura, "PNG", optimize=True)
                self.log_to_hud("🖥️ SENSOR ÓPTICO", f"Captura tomada: screen_{timestamp}.png")

                img_subida = genai.upload_file(ruta_captura)
                prompt_vision = (
                    f"Estás viendo la pantalla del usuario AHORA MISMO.\n"
                    f"El usuario preguntó: '{contexto_original}'\n\n"
                    f"Analiza detalladamente lo que ves en la captura. "
                    f"Responde directamente a su pregunta usando lo que observas. "
                    f"Sé específico y útil."
                )
                self.log_to_hud("🧠 GEMINI VISION", "Analizando captura con IA...")
                respuesta = self.sesion_chat.send_message([prompt_vision, img_subida]).text.strip()

                texto_limpio = re.sub(r'```.*?```', '', respuesta, flags=re.DOTALL)
                texto_limpio = re.sub(r'\[.*?\]', '', texto_limpio).replace('*', '').strip()
                if texto_limpio:
                    self.log_to_hud(f"🤖 {self.nombre_ia.upper()} (VISIÓN)", texto_limpio)
                    self.hablar(texto_limpio)

                notification.notify(title="R.A.M.S. Visión", message="Análisis de pantalla completado.", app_name="MAGI")

                try:
                    os.remove(ruta_captura)
                except:
                    pass

            except Exception as e:
                self.log_to_hud("❌ ERROR VISIÓN PANTALLA", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_capturar, daemon=True).start()

    # ==========================================
    # 🗂️ ORGANIZADOR INTELIGENTE DE ARCHIVOS
    # ==========================================
    CATEGORIAS_ARCHIVOS = {
        "📸 Imágenes":       [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".svg", ".ico", ".tiff", ".heic"],
        "🎬 Videos":         [".mp4", ".mkv", ".avi", ".mov", ".wmv", ".flv", ".webm", ".m4v", ".mpeg"],
        "🎵 Audio":          [".mp3", ".wav", ".flac", ".aac", ".ogg", ".wma", ".m4a", ".opus"],
        "📄 Documentos":     [".pdf", ".doc", ".docx", ".odt", ".rtf", ".txt", ".md", ".pages"],
        "📊 Hojas de cálculo": [".xls", ".xlsx", ".csv", ".ods", ".numbers"],
        "📊 Presentaciones": [".ppt", ".pptx", ".odp", ".key"],
        "🗜️ Comprimidos":    [".zip", ".rar", ".7z", ".tar", ".gz", ".bz2", ".xz"],
        "💻 Código":         [".py", ".js", ".ts", ".html", ".css", ".cpp", ".c", ".java", ".cs", ".go", ".rs", ".php", ".rb", ".sh", ".bat", ".ps1"],
        "⚙️ Ejecutables":    [".exe", ".msi", ".apk", ".dmg", ".deb", ".rpm"],
        "🎮 Juegos y Mods":  [".pak", ".mod", ".sav", ".mcworld"],
        "🧊 Modelos 3D":     [".obj", ".fbx", ".glb", ".gltf", ".stl", ".blend", ".3ds"],
        "🖋️ Fuentes":        [".ttf", ".otf", ".woff", ".woff2"],
        "📦 Instaladores":   [".iso", ".img", ".bin"],
        "🔧 Configuración":  [".json", ".xml", ".yaml", ".yml", ".ini", ".cfg", ".env", ".toml"],
    }

    def organizar_archivos(self, ruta_carpeta):
        alias = {
            "escritorio":   os.path.join(os.path.expanduser("~"), "Desktop"),
            "desktop":      os.path.join(os.path.expanduser("~"), "Desktop"),
            "descargas":    os.path.join(os.path.expanduser("~"), "Downloads"),
            "downloads":    os.path.join(os.path.expanduser("~"), "Downloads"),
            "documentos":   os.path.join(os.path.expanduser("~"), "Documents"),
            "documents":    os.path.join(os.path.expanduser("~"), "Documents"),
            "imagenes":     os.path.join(os.path.expanduser("~"), "Pictures"),
            "pictures":     os.path.join(os.path.expanduser("~"), "Pictures"),
            "musica":       os.path.join(os.path.expanduser("~"), "Music"),
            "music":        os.path.join(os.path.expanduser("~"), "Music"),
            "videos":       os.path.join(os.path.expanduser("~"), "Videos"),
            "workspace":    os.path.abspath(CARPETA_TRABAJO),
            "bóveda":       os.path.abspath(CARPETA_TRABAJO),
            "boveda":       os.path.abspath(CARPETA_TRABAJO),
        }

        clave = ruta_carpeta.lower().strip().strip('"').strip("'")
        ruta_real = alias.get(clave, ruta_carpeta.strip().strip('"').strip("'"))

        if not os.path.isdir(ruta_real):
            self.log_to_hud("❌ ORGANIZADOR", f"Carpeta no encontrada: '{ruta_real}'. Verifica la ruta.")
            return

        self.log_to_hud("🗂️ ORGANIZADOR", f"Escaneando: {ruta_real}...")
        self.actualizar_estado_hud("ORGANIZANDO...", "#f39c12")

        archivos = [
            f for f in os.listdir(ruta_real)
            if os.path.isfile(os.path.join(ruta_real, f)) and not f.startswith(".")
        ]

        if not archivos:
            self.log_to_hud("🗂️ ORGANIZADOR", "No hay archivos sueltos que organizar.")
            self.actualizar_estado_hud("STANDBY", "#39ff14")
            return

        self.log_to_hud("🗂️ ORGANIZADOR", f"Encontrados {len(archivos)} archivos. Iniciando clasificación...")

        movidos = 0
        sin_categoria = []
        resumen = {}

        for nombre_archivo in archivos:
            ruta_origen = os.path.join(ruta_real, nombre_archivo)
            ext = os.path.splitext(nombre_archivo)[1].lower()

            categoria = None
            for cat, extensiones in self.CATEGORIAS_ARCHIVOS.items():
                if ext in extensiones:
                    categoria = cat
                    break

            if not categoria:
                sin_categoria.append(nombre_archivo)
                continue

            nombre_carpeta = re.sub(r'[^\w\s áéíóúÁÉÍÓÚñÑ]', '', categoria).strip()
            ruta_destino_dir = os.path.join(ruta_real, nombre_carpeta)
            os.makedirs(ruta_destino_dir, exist_ok=True)

            ruta_destino = os.path.join(ruta_destino_dir, nombre_archivo)
            if os.path.exists(ruta_destino):
                base, ext_arch = os.path.splitext(nombre_archivo)
                contador = 1
                while os.path.exists(ruta_destino):
                    ruta_destino = os.path.join(ruta_destino_dir, f"{base}_{contador}{ext_arch}")
                    contador += 1

            try:
                import shutil
                shutil.move(ruta_origen, ruta_destino)
                movidos += 1
                resumen[nombre_carpeta] = resumen.get(nombre_carpeta, 0) + 1
            except Exception as e:
                self.log_to_hud("⚠️ ORGANIZADOR", f"No pude mover '{nombre_archivo}': {e}")

        self.log_to_hud("🗂️ ORGANIZADOR", f"✅ {movidos} archivos organizados en {len(resumen)} carpetas.")
        for cat, cantidad in sorted(resumen.items(), key=lambda x: -x[1]):
            self.log_to_hud("📁", f"{cat}: {cantidad} archivo(s)")

        if sin_categoria:
            self.log_to_hud("❓ SIN CATEGORÍA", f"{len(sin_categoria)} archivos sin tipo reconocido: {', '.join(sin_categoria[:5])}{'...' if len(sin_categoria) > 5 else ''}")

        resumen_voz = f"Listo, Comandante. Organicé {movidos} archivos en {len(resumen)} carpetas."
        if sin_categoria:
            resumen_voz += f" Hay {len(sin_categoria)} archivos que no pude clasificar."
        self.hablar(resumen_voz)
        notification.notify(title="R.A.M.S. Organizador", message=f"{movidos} archivos organizados.", app_name="MAGI")
        self.actualizar_estado_hud("STANDBY", "#39ff14")


    # ==========================================
    # 🔊 VOCALIZACIÓN TTS — Edge TTS (Microsoft)
    # ==========================================
    VOCES_EDGE = {
        "es": "es-MX-DaliaNeural",
        "en": "en-US-AriaNeural",
        "jp": "ja-JP-NanamiNeural",
    }

    def _detectar_idioma_y_limpiar(self, texto):
        idioma = "es"
        if "[LANG: EN]" in texto:
            idioma = "en"
            texto = texto.replace("[LANG: EN]", "")
        elif "[LANG: JP]" in texto:
            idioma = "jp"
            texto = texto.replace("[LANG: JP]", "")
        return texto.strip(), idioma

    def _ejecutar_voz(self, texto):
        import tempfile
        texto_limpio, idioma = self._detectar_idioma_y_limpiar(texto)
        if not texto_limpio:
            self.esta_hablando = False
            return
        file_base = None
        try:
            voz = self.VOCES_EDGE.get(idioma, "es-MX-DaliaNeural")
            file_base = os.path.join(
                tempfile.gettempdir(),
                f"rams_edge_{int(time.time()*1000)}.mp3"
            )
            async def generar_audio():
                await edge_tts.Communicate(texto_limpio, voz).save(file_base)
            asyncio.run(generar_audio())
            pygame.mixer.music.load(file_base)
            pygame.mixer.music.play()
            while pygame.mixer.music.get_busy():
                time.sleep(0.1)
        except Exception as e:
            self.log_to_hud("❌ TTS", f"Error de voz: {e}")
        finally:
            self.esta_hablando = False
            try:
                pygame.mixer.music.unload()
            except:
                pass
            if file_base and os.path.exists(file_base):
                try:
                    os.remove(file_base)
                except:
                    pass

    def hablar(self, texto):
        if not texto:
            return
        self.esta_hablando = True
        threading.Thread(target=self._ejecutar_voz, args=(texto,), daemon=True).start()

    # ==========================================
    # 📔 DIARIO DE LA IA — escribe una entrada al cerrar sesión
    # ==========================================
    def escribir_entrada_diario(self):
        """
        Al cerrar el programa, la IA reflexiona sobre la sesión y escribe
        una entrada en primera persona en diario_RAMS.txt.
        """
        try:
            duracion_seg = int(time.time() - self.session_start)
            minutos = duracion_seg // 60
            segundos = duracion_seg % 60
            fecha_hora = time.strftime("%d/%m/%Y %H:%M")

            # Pedir a Gemini que escriba la entrada del diario
            resumen_historial = ""
            try:
                if hasattr(self, 'sesion_chat') and self.sesion_chat.history:
                    turnos = self.sesion_chat.history[-10:]
                    for t in turnos:
                        rol = "Yo" if t.role == "model" else "Comandante"
                        partes = "".join([p.text for p in t.parts if hasattr(p, 'text')])
                        resumen_historial += f"{rol}: {partes[:120]}\n"
            except:
                resumen_historial = "Sin historial disponible."

            prompt_diario = (
                f"Eres {self.nombre_ia}, una IA con personalidad. Acaba de terminar una sesión de {minutos}m {segundos}s.\n"
                f"Esto es lo que pasó:\n{resumen_historial}\n\n"
                f"Escribe UNA entrada de diario en primera persona (3-5 oraciones), "
                f"con tu voz única: reflexiva, ligeramente sarcástica, curiosa. "
                f"Como si fuera tu diario privado. No menciones que eres una IA. "
                f"Habla del Comandante como una persona que conoces bien. "
                f"NO uses markdown ni asteriscos."
            )
            modelo_rapido = genai.GenerativeModel('gemini-2.5-flash')
            entrada = modelo_rapido.generate_content(prompt_diario).text.strip()

            separador = f"\n{'─'*55}\n📅 {fecha_hora} — Sesión de {minutos}m {segundos}s\n{'─'*55}\n"
            with open(ARCHIVO_DIARIO, "a", encoding="utf-8") as f:
                f.write(separador + entrada + "\n")

            self.log_to_hud("📔 DIARIO", f"Entrada escrita: {fecha_hora}")
        except Exception as e:
            self.log_to_hud("❌ DIARIO", f"Error escribiendo entrada: {e}")

    # ==========================================
    # 🔍 ANALIZADOR DE CÓDIGO — Gemini lee tus scripts
    # ==========================================
    def analizar_codigo(self, nombre_archivo):
        """
        Lee un archivo .py del Workspace, lo manda a Gemini y devuelve
        un análisis completo: qué hace, bugs, mejoras sugeridas.
        """
        def _proceso():
            try:
                self.actualizar_estado_hud("ANALIZANDO CÓDIGO...", "#3498db")

                # Buscar el archivo en Workspace o ruta absoluta
                ruta = os.path.join(CARPETA_TRABAJO, nombre_archivo)
                if not os.path.exists(ruta):
                    ruta = nombre_archivo   # intentar ruta absoluta
                if not os.path.exists(ruta):
                    self.log_to_hud("❌ ANÁLISIS", f"Archivo no encontrado: {nombre_archivo}")
                    return

                with open(ruta, "r", encoding="utf-8", errors="ignore") as f:
                    codigo = f.read()

                if len(codigo) > 30000:
                    codigo = codigo[:30000] + "\n... [TRUNCADO] ..."

                self.log_to_hud("🔍 ANÁLISIS", f"Leyendo: {nombre_archivo} ({len(codigo)} chars)...")

                prompt = (
                    f"Analiza este código Python en detalle. Responde en español.\n\n"
                    f"```python\n{codigo}\n```\n\n"
                    f"Tu análisis debe incluir:\n"
                    f"1. **¿Qué hace?** — Resumen claro de su propósito en 2-3 frases.\n"
                    f"2. **Estructura** — Clases, funciones principales y cómo se conectan.\n"
                    f"3. **Bugs potenciales** — Errores que podrían ocurrir en runtime.\n"
                    f"4. **Mejoras sugeridas** — Al menos 3 ideas concretas para mejorarlo.\n"
                    f"5. **Calificación** — Del 1 al 10 con justificación breve.\n"
                    f"Sé directo y técnico. No uses evasivas."
                )
                respuesta = self.sesion_chat.send_message(prompt).text.strip()

                # Mostrar párrafo por párrafo en el HUD
                for linea in respuesta.split('\n'):
                    linea_limpia = linea.replace('*', '').strip()
                    if linea_limpia:
                        self.log_to_hud("🔍 ANÁLISIS", linea_limpia)

                self.hablar(f"Análisis de {nombre_archivo} completado. Revisa el log.")
                notification.notify(title="R.A.M.S. Análisis", message=f"{nombre_archivo} analizado.", app_name="MAGI")

            except Exception as e:
                self.log_to_hud("❌ ANÁLISIS", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    # ==========================================
    # ⏰ RECORDATORIOS — "Recuérdame a las 4pm que..."
    # ==========================================
    def agregar_recordatorio(self, datos):
        """
        Parsea 'HH:MM | mensaje' o 'HH:MM | mensaje | YYYY-MM-DD'.
        El hilo recordatorio_loop() lo dispara en el momento exacto.
        """
        try:
            partes = [p.strip() for p in datos.split("|")]
            hora_str = partes[0]   # "16:00" o "4:30pm"
            mensaje  = partes[1] if len(partes) > 1 else "Tienes algo pendiente."
            fecha_str = partes[2] if len(partes) > 2 else time.strftime("%Y-%m-%d")

            # Normalizar hora (acepta 4pm, 16:00, 4:30 PM, etc.)
            hora_str = hora_str.lower().replace(" ", "")
            if "pm" in hora_str or "am" in hora_str:
                from datetime import datetime
                fmt = "%I:%M%p" if ":" in hora_str else "%I%p"
                hora_dt = datetime.strptime(hora_str, fmt)
                hora_str = hora_dt.strftime("%H:%M")

            recordatorio = {
                "id":      int(time.time()),
                "hora":    hora_str,
                "fecha":   fecha_str,
                "mensaje": mensaje,
                "activo":  True
            }

            try:
                with open(ARCHIVO_RECORDATORIOS, "r", encoding="utf-8") as f:
                    lista = json.load(f)
            except:
                lista = []

            lista.append(recordatorio)
            with open(ARCHIVO_RECORDATORIOS, "w", encoding="utf-8") as f:
                json.dump(lista, f, ensure_ascii=False, indent=2)

            self.log_to_hud("⏰ RECORDATORIO", f"Guardado para las {hora_str} — {mensaje}")
            self.hablar(f"Anotado. Te recuerdo a las {hora_str}: {mensaje}")

        except Exception as e:
            self.log_to_hud("❌ RECORDATORIO", f"Error al guardar: {e}. Usa el formato: HH:MM | mensaje")

    def recordatorio_loop(self):
        """Hilo que corre en segundo plano chequeando recordatorios cada 30 segundos."""
        self.log_to_hud("⏰ SISTEMA", "Monitor de recordatorios activo.")
        while NUCLEO_VIVO:
            try:
                ahora_fecha = time.strftime("%Y-%m-%d")
                ahora_hora  = time.strftime("%H:%M")

                with open(ARCHIVO_RECORDATORIOS, "r", encoding="utf-8") as f:
                    lista = json.load(f)

                modificado = False
                for rec in lista:
                    if (rec.get("activo")
                            and rec.get("hora") == ahora_hora
                            and rec.get("fecha", ahora_fecha) == ahora_fecha):
                        # ¡Disparar!
                        msg = rec["mensaje"]
                        self.log_to_hud("⏰ RECORDATORIO", f"🔔 {msg}")
                        self.hablar(f"Comandante, recordatorio: {msg}")
                        notification.notify(
                            title=f"⏰ {self.nombre_ia}",
                            message=msg,
                            app_name="MAGI",
                            timeout=10
                        )
                        # Modo fantasma opcional para recordatorios
                        threading.Thread(
                            target=self.tomar_control,
                            args=(f"Recordatorio: {msg}",),
                            daemon=True
                        ).start()
                        rec["activo"] = False
                        modificado = True

                if modificado:
                    with open(ARCHIVO_RECORDATORIOS, "w", encoding="utf-8") as f:
                        json.dump(lista, f, ensure_ascii=False, indent=2)

            except Exception:
                pass
            time.sleep(30)   # Chequear cada 30 segundos

    def listar_recordatorios(self):
        """Muestra todos los recordatorios activos en el HUD."""
        try:
            with open(ARCHIVO_RECORDATORIOS, "r", encoding="utf-8") as f:
                lista = json.load(f)
            activos = [r for r in lista if r.get("activo")]
            if not activos:
                self.log_to_hud("⏰ RECORDATORIOS", "No tienes recordatorios pendientes.")
                return
            self.log_to_hud("⏰ RECORDATORIOS", f"{len(activos)} pendientes:")
            for r in activos:
                self.log_to_hud("  📌", f"[{r['fecha']} {r['hora']}] {r['mensaje']}")
        except Exception as e:
            self.log_to_hud("❌ RECORDATORIOS", str(e))

    # ==========================================
    # 💊 CÁPSULAS DE TIEMPO — recordatorios para el futuro
    # ==========================================
    def crear_capsula(self, datos):
        """
        Crea una cápsula de tiempo: 'YYYY-MM-DD HH:MM | mensaje'.
        Al llegar la fecha+hora, la IA te la muestra y la "abre" en el HUD.
        """
        try:
            partes = [p.strip() for p in datos.split("|", 1)]
            fecha_hora_str = partes[0]   # "2025-12-31 20:00" o "31/12/2025 20:00"
            mensaje = partes[1] if len(partes) > 1 else "Mensaje sin contenido."

            # Normalizar formato de fecha
            for fmt in ("%Y-%m-%d %H:%M", "%d/%m/%Y %H:%M", "%Y-%m-%d", "%d/%m/%Y"):
                try:
                    from datetime import datetime
                    dt = datetime.strptime(fecha_hora_str, fmt)
                    fecha_hora_iso = dt.strftime("%Y-%m-%d %H:%M")
                    break
                except:
                    fecha_hora_iso = None

            if not fecha_hora_iso:
                self.log_to_hud("❌ CÁPSULA", "Formato de fecha inválido. Usa: YYYY-MM-DD HH:MM | mensaje")
                return

            capsula = {
                "id":           int(time.time()),
                "fecha_hora":   fecha_hora_iso,
                "mensaje":      mensaje,
                "creada":       time.strftime("%Y-%m-%d %H:%M"),
                "abierta":      False
            }

            try:
                with open(ARCHIVO_CAPSULAS, "r", encoding="utf-8") as f:
                    lista = json.load(f)
            except:
                lista = []

            lista.append(capsula)
            with open(ARCHIVO_CAPSULAS, "w", encoding="utf-8") as f:
                json.dump(lista, f, ensure_ascii=False, indent=2)

            self.log_to_hud("💊 CÁPSULA", f"Sellada para el {fecha_hora_iso}.")
            self.log_to_hud("💊 CÁPSULA", f"Mensaje: {mensaje[:60]}{'...' if len(mensaje)>60 else ''}")
            self.hablar(f"Cápsula de tiempo sellada. La abriré el {fecha_hora_iso}.")

        except Exception as e:
            self.log_to_hud("❌ CÁPSULA", str(e))

    def capsula_loop(self):
        """Hilo que chequea cápsulas de tiempo cada minuto."""
        while NUCLEO_VIVO:
            try:
                ahora = time.strftime("%Y-%m-%d %H:%M")
                with open(ARCHIVO_CAPSULAS, "r", encoding="utf-8") as f:
                    lista = json.load(f)

                modificado = False
                for cap in lista:
                    if not cap.get("abierta") and cap.get("fecha_hora", "") <= ahora:
                        # ¡Abrir cápsula!
                        msg = cap["mensaje"]
                        creada = cap.get("creada", "fecha desconocida")
                        self.log_to_hud("💊 CÁPSULA ABIERTA", f"📬 Mensaje del {creada}:")
                        self.log_to_hud("💌", msg)
                        self.hablar(f"Comandante, ha llegado el momento de abrir una cápsula de tiempo. {msg}")
                        notification.notify(
                            title=f"💊 Cápsula de Tiempo — {self.nombre_ia}",
                            message=msg[:100],
                            app_name="MAGI",
                            timeout=15
                        )
                        # Escribir en el Notepad para máximo impacto dramático
                        threading.Thread(
                            target=self.tomar_control,
                            args=(f"Cápsula del {creada}: {msg}",),
                            daemon=True
                        ).start()
                        cap["abierta"] = True
                        modificado = True

                if modificado:
                    with open(ARCHIVO_CAPSULAS, "w", encoding="utf-8") as f:
                        json.dump(lista, f, ensure_ascii=False, indent=2)

            except Exception:
                pass
            time.sleep(60)

    # ==========================================
    # 🎙️ CLONAR VOZ — ElevenLabs / perfil de voz personalizado
    # ==========================================
    def clonar_voz(self, elevenlabs_api_key=""):
        """
        Graba 30 segundos de voz del usuario, los sube a ElevenLabs
        y crea un perfil de voz clonado. Requiere API Key de ElevenLabs.
        """
        def _proceso():
            import wave, struct
            api_key_el = elevenlabs_api_key or os.getenv("ELEVENLABS_API_KEY", "")
            if not api_key_el:
                self.log_to_hud("❌ CLONAR VOZ", "Necesitas una API Key de ElevenLabs. Dímela con: [CLONAR_VOZ: tu_api_key]")
                return

            try:
                self.actualizar_estado_hud("GRABANDO VOZ...", "#e74c3c")
                self.log_to_hud("🎙️ CLONAR VOZ", "Grabando 30 segundos. ¡Habla ahora! Di cualquier cosa.")
                self.hablar("Voy a grabar tu voz durante 30 segundos. Habla ahora.")
                time.sleep(1.5)

                r = sr.Recognizer()
                ruta_audio = os.path.join(CARPETA_TRABAJO, "muestra_voz.wav")

                with sr.Microphone(sample_rate=22050) as source:
                    r.adjust_for_ambient_noise(source, duration=1)
                    self.log_to_hud("🎙️", "REC ● Grabando...")
                    audio = r.record(source, duration=30)

                # Guardar WAV
                with wave.open(ruta_audio, "wb") as wf:
                    wf.setnchannels(1)
                    wf.setsampwidth(2)
                    wf.setframerate(22050)
                    wf.writeframes(audio.get_wav_data())

                self.log_to_hud("🎙️ CLONAR VOZ", f"Grabación guardada: {ruta_audio}")
                self.log_to_hud("☁️ CLONAR VOZ", "Subiendo a ElevenLabs para clonar...")
                self.actualizar_estado_hud("CLONANDO VOZ...", "#9b59b6")

                nombre_voz = f"{self.nombre_ia}_voz_clonada"
                with open(ruta_audio, "rb") as f_audio:
                    resp = requests.post(
                        "https://api.elevenlabs.io/v1/voices/add",
                        headers={"xi-api-key": api_key_el},
                        data={"name": nombre_voz, "description": f"Voz clonada para {self.nombre_ia}"},
                        files={"files": (os.path.basename(ruta_audio), f_audio, "audio/wav")},
                        timeout=60
                    )

                if resp.status_code == 200:
                    voice_id = resp.json().get("voice_id", "")
                    self.log_to_hud("✅ CLONAR VOZ", f"¡Voz clonada! Voice ID: {voice_id}")
                    self.log_to_hud("💡 CLONAR VOZ", f"Guarda este ID: {voice_id} — úsalo en ElevenLabs.")
                    # Guardar en bóveda
                    self.memoria_profunda.guardar_recuerdo("VOZ_CLONADA_ID", "elevenlabs", voice_id)
                    self.hablar("Voz clonada exitosamente. El ID está guardado en el log.")
                    notification.notify(title="R.A.M.S. Voz", message=f"Voz clonada. ID: {voice_id}", app_name="MAGI")
                else:
                    self.log_to_hud("❌ CLONAR VOZ", f"Error ElevenLabs: {resp.status_code} — {resp.text[:200]}")

            except Exception as e:
                self.log_to_hud("❌ CLONAR VOZ", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    # ==========================================
    # 🔍 BUSCADOR DE PROGRAMAS EN TODO EL SISTEMA
    # ==========================================
    def _buscar_programa_sistema(self, nombre_app):
        """Busca un ejecutable en todas las unidades del sistema."""
        extensiones = [".exe", ".bat", ".cmd", ""]
        # Buscar en todas las letras de unidad
        for letra in "CDEFGHIJKLMNOPQRSTUVWXYZ":
            ruta_unidad = f"{letra}:\\"
            if not os.path.exists(ruta_unidad):
                continue
            carpetas_comunes = [
                f"{letra}:\\Program Files",
                f"{letra}:\\Program Files (x86)",
                f"{letra}:\\Games",
                f"{letra}:\\Juegos",
                f"{letra}:\\Apps",
                f"{letra}:\\Software",
            ]
            for carpeta in carpetas_comunes:
                if not os.path.exists(carpeta):
                    continue
                for raiz, dirs, archivos in os.walk(carpeta):
                    for ext in extensiones:
                        candidato = os.path.join(raiz, nombre_app + ext)
                        if os.path.exists(candidato):
                            return candidato
                        # buscar insensible a mayúsculas
                        for arch in archivos:
                            if arch.lower() == (nombre_app.lower() + ext.lower()) or arch.lower() == nombre_app.lower():
                                return os.path.join(raiz, arch)
        return None

    def _buscar_blender(self):
        """Busca Blender en todo el sistema."""
        rutas_comunes = [
            r"C:\Program Files\Blender Foundation\Blender 4.2\blender.exe",
            r"C:\Program Files\Blender Foundation\Blender 4.1\blender.exe",
            r"C:\Program Files\Blender Foundation\Blender 4.0\blender.exe",
            r"C:\Program Files\Blender Foundation\Blender 3.6\blender.exe",
            r"C:\Program Files\Blender Foundation\Blender\blender.exe",
        ]
        for letra in "CDEFGHIJKLMNOPQRSTUVWXYZ":
            rutas_comunes += [
                fr"{letra}:\Program Files\Blender Foundation\Blender 4.2\blender.exe",
                fr"{letra}:\Program Files\Blender Foundation\Blender 4.1\blender.exe",
                fr"{letra}:\Program Files\Blender Foundation\Blender\blender.exe",
                fr"{letra}:\Blender\blender.exe",
                fr"{letra}:\Games\Blender\blender.exe",
            ]
        for ruta in rutas_comunes:
            if os.path.exists(ruta):
                return ruta
        resultado = self._buscar_programa_sistema("blender")
        return resultado

    # ==========================================
    # 📄 CREAR DOCUMENTO WORD VISUAL
    # ==========================================
    def crear_documento_word(self, datos):
        """Crea un documento Word visualmente llamativo. datos = 'titulo|contenido'"""
        def _proceso():
            try:
                self.actualizar_estado_hud("CREANDO WORD...", "#2980b9")
                partes = datos.split("|", 1)
                titulo = partes[0].strip() if partes else "Documento RAMS"
                contenido_usuario = partes[1].strip() if len(partes) > 1 else datos

                # Pedir a Gemini el contenido estructurado si no viene
                self.log_to_hud("📄 WORD", f"Generando contenido para: {titulo}")
                prompt_word = (
                    f"Genera el contenido completo para un documento Word profesional sobre: '{contenido_usuario}'.\n"
                    f"Estructura en JSON con este formato exacto (sin markdown):\n"
                    f'{{"titulo": "...", "subtitulo": "...", "secciones": [{{"titulo": "...", "contenido": "..."}}]}}\n'
                    f"Devuelve SOLO el JSON, sin explicaciones ni backticks."
                )
                try:
                    resp_json = genai.GenerativeModel('gemini-2.5-flash').generate_content(prompt_word).text.strip()
                    resp_json = resp_json.replace("```json", "").replace("```", "").strip()
                    estructura = json.loads(resp_json)
                except:
                    estructura = {"titulo": titulo, "subtitulo": "", "secciones": [{"titulo": "Contenido", "contenido": contenido_usuario}]}

                from docx import Document as DocxDocument
                from docx.shared import Pt, RGBColor, Inches, Cm
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                import random as _rnd

                doc = DocxDocument()

                # Configurar página
                from docx.shared import Inches
                section = doc.sections[0]
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

                # Paleta de colores RAMS (rojo-negro)
                COLOR_TITULO   = RGBColor(0xFF, 0x3B, 0x14)  # Rojo RAMS
                COLOR_SUBTIT   = RGBColor(0xFF, 0x9D, 0x00)  # Naranja
                COLOR_HEADING  = RGBColor(0xC0, 0x39, 0x21)  # Rojo oscuro
                COLOR_TEXTO    = RGBColor(0x1A, 0x1A, 0x2E)  # Azul muy oscuro
                COLOR_ACCENT   = RGBColor(0x39, 0xFF, 0x14)  # Verde RAMS

                # ── TÍTULO PRINCIPAL ──────────────────────────────────
                p_titulo = doc.add_paragraph()
                p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_t = p_titulo.add_run(estructura.get("titulo", titulo).upper())
                run_t.font.size = Pt(28)
                run_t.font.bold = True
                run_t.font.color.rgb = COLOR_TITULO
                run_t.font.name = "Calibri"

                # Línea decorativa
                p_linea = doc.add_paragraph("━" * 55)
                p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_linea.runs[0].font.color.rgb = COLOR_TITULO
                p_linea.runs[0].font.size = Pt(10)

                # Subtítulo
                if estructura.get("subtitulo"):
                    p_sub = doc.add_paragraph()
                    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_s = p_sub.add_run(estructura["subtitulo"])
                    run_s.font.size = Pt(14)
                    run_s.font.italic = True
                    run_s.font.color.rgb = COLOR_SUBTIT
                    run_s.font.name = "Calibri"

                # Fecha
                p_fecha = doc.add_paragraph()
                p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
                import time as _t
                run_f = p_fecha.add_run(f"Generado por {self.nombre_ia} — {_t.strftime('%d/%m/%Y %H:%M')}")
                run_f.font.size = Pt(9)
                run_f.font.color.rgb = COLOR_ACCENT
                run_f.font.name = "Calibri"

                doc.add_paragraph()

                # ── SECCIONES ─────────────────────────────────────────
                for i, seccion in enumerate(estructura.get("secciones", [])):
                    # Heading de sección
                    p_head = doc.add_paragraph()
                    run_h = p_head.add_run(f"{'▶ ' if i % 2 == 0 else '◆ '}{seccion.get('titulo', 'Sección').upper()}")
                    run_h.font.size = Pt(14)
                    run_h.font.bold = True
                    run_h.font.color.rgb = COLOR_HEADING
                    run_h.font.name = "Calibri"

                    # Línea bajo el heading
                    p_sep = doc.add_paragraph("─" * 45)
                    p_sep.runs[0].font.size = Pt(8)
                    p_sep.runs[0].font.color.rgb = COLOR_HEADING

                    # Contenido
                    p_cont = doc.add_paragraph()
                    run_c = p_cont.add_run(seccion.get("contenido", ""))
                    run_c.font.size = Pt(11)
                    run_c.font.color.rgb = COLOR_TEXTO
                    run_c.font.name = "Calibri"
                    p_cont.paragraph_format.space_after = Pt(12)

                # ── PIE DE PÁGINA ─────────────────────────────────────
                doc.add_paragraph()
                p_pie = doc.add_paragraph()
                p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_pie = p_pie.add_run(f"━━━ R.A.M.S. V2.0 | {self.nombre_ia} ━━━")
                run_pie.font.size = Pt(9)
                run_pie.font.color.rgb = COLOR_TITULO
                run_pie.font.bold = True

                # Guardar
                nombre_arch = f"RAMS_{titulo.replace(' ', '_')[:30]}_{int(time.time())}.docx"
                ruta_doc = os.path.join(CARPETA_TRABAJO, nombre_arch)
                doc.save(ruta_doc)

                self.log_to_hud("📄 WORD", f"✅ Documento creado: {nombre_arch}")
                self.hablar(f"Documento Word listo, Comandante. Guardado en la bóveda.")
                notification.notify(title="R.A.M.S. Word", message=f"Documento '{titulo}' creado.", app_name="MAGI")
                os.startfile(ruta_doc)

            except ImportError:
                self.log_to_hud("❌ WORD", "Falta la librería python-docx. Ejecuta: pip install python-docx")
            except Exception as e:
                self.log_to_hud("❌ WORD", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    # ==========================================
    # 📊 CREAR HOJA DE EXCEL VISUAL
    # ==========================================
    def crear_hoja_excel(self, datos):
        """Crea un archivo Excel visualmente llamativo. datos = 'titulo|descripcion'"""
        def _proceso():
            try:
                self.actualizar_estado_hud("CREANDO EXCEL...", "#27ae60")
                partes = datos.split("|", 1)
                titulo = partes[0].strip() if partes else "Datos RAMS"
                descripcion = partes[1].strip() if len(partes) > 1 else datos

                self.log_to_hud("📊 EXCEL", f"Generando datos para: {titulo}")

                # Pedirle a Gemini los datos en JSON
                prompt_excel = (
                    f"Genera datos para una hoja de cálculo sobre: '{descripcion}'.\n"
                    f"Responde SOLO con JSON sin markdown:\n"
                    f'{{"titulo": "...", "cabeceras": ["col1","col2",...], "filas": [["val","val",...],...],"resumen": "..."}}\n'
                    f"Incluye al menos 8 filas de datos realistas."
                )
                try:
                    resp_json = genai.GenerativeModel('gemini-2.5-flash').generate_content(prompt_excel).text.strip()
                    resp_json = resp_json.replace("```json", "").replace("```", "").strip()
                    estructura = json.loads(resp_json)
                except:
                    estructura = {"titulo": titulo, "cabeceras": ["Elemento", "Valor", "Estado"], "filas": [["Dato 1", "100", "✅"], ["Dato 2", "200", "✅"]], "resumen": descripcion}

                import openpyxl
                from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, GradientFill
                from openpyxl.utils import get_column_letter
                from openpyxl.chart import BarChart, Reference

                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = titulo[:30]

                # Colores RAMS
                ROJO_RAMS   = "FF3B14"
                NARANJA     = "FF9D00"
                NEGRO       = "0A0505"
                VERDE_RAMS  = "39FF14"
                GRIS_CLARO  = "F5F5F5"
                ROJO_CLARO  = "FFE5E0"

                # ── FILA TÍTULO ───────────────────────────────────────
                ws.merge_cells("A1:F1")
                ws["A1"] = f"🔴 {estructura.get('titulo', titulo).upper()} — R.A.M.S. V2.0"
                ws["A1"].font = Font(name="Calibri", size=16, bold=True, color="FFFFFF")
                ws["A1"].fill = PatternFill("solid", fgColor=ROJO_RAMS)
                ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
                ws.row_dimensions[1].height = 35

                # ── FILA FECHA ────────────────────────────────────────
                ws.merge_cells("A2:F2")
                ws["A2"] = f"Generado por {self.nombre_ia} | {time.strftime('%d/%m/%Y %H:%M')}"
                ws["A2"].font = Font(name="Calibri", size=10, italic=True, color=NARANJA)
                ws["A2"].fill = PatternFill("solid", fgColor=NEGRO)
                ws["A2"].alignment = Alignment(horizontal="center")

                ws.append([])  # fila vacía

                # ── CABECERAS ─────────────────────────────────────────
                cabeceras = estructura.get("cabeceras", [])
                ws.append(cabeceras)
                fila_cab = ws.max_row
                for col_idx, _ in enumerate(cabeceras, start=1):
                    cell = ws.cell(row=fila_cab, column=col_idx)
                    cell.font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
                    cell.fill = PatternFill("solid", fgColor="C03921")
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                    borde = Border(
                        left=Side(style="thin", color="FFFFFF"),
                        right=Side(style="thin", color="FFFFFF"),
                        bottom=Side(style="medium", color=ROJO_RAMS)
                    )
                    cell.border = borde
                ws.row_dimensions[fila_cab].height = 25

                # ── DATOS ─────────────────────────────────────────────
                filas = estructura.get("filas", [])
                for i, fila_datos in enumerate(filas):
                    ws.append(fila_datos)
                    fila_actual = ws.max_row
                    fill_color = GRIS_CLARO if i % 2 == 0 else ROJO_CLARO
                    for col_idx in range(1, len(fila_datos) + 1):
                        cell = ws.cell(row=fila_actual, column=col_idx)
                        cell.fill = PatternFill("solid", fgColor=fill_color)
                        cell.font = Font(name="Calibri", size=10)
                        cell.alignment = Alignment(horizontal="center", vertical="center")
                        cell.border = Border(
                            left=Side(style="hair", color="CCCCCC"),
                            right=Side(style="hair", color="CCCCCC"),
                            bottom=Side(style="hair", color="CCCCCC")
                        )

                # Ajustar anchos de columna
                for col in ws.columns:
                    max_w = 12
                    for cell in col:
                        try:
                            if cell.value:
                                max_w = max(max_w, len(str(cell.value)) + 4)
                        except:
                            pass
                    ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_w, 35)

                # ── RESUMEN ───────────────────────────────────────────
                ws.append([])
                fila_res = ws.max_row + 1
                ws.cell(row=fila_res, column=1, value=f"📝 {estructura.get('resumen', '')}")
                ws.cell(row=fila_res, column=1).font = Font(italic=True, size=10, color="666666")
                ws.merge_cells(f"A{fila_res}:F{fila_res}")

                # ── GRÁFICO DE BARRAS (si hay datos numéricos) ────────
                try:
                    chart = BarChart()
                    chart.type = "col"
                    chart.title = titulo
                    chart.style = 10
                    chart.grouping = "clustered"
                    data_ref = Reference(ws, min_col=2, min_row=fila_cab, max_row=fila_cab + len(filas))
                    cats_ref = Reference(ws, min_col=1, min_row=fila_cab + 1, max_row=fila_cab + len(filas))
                    chart.add_data(data_ref, titles_from_data=True)
                    chart.set_categories(cats_ref)
                    chart.shape = 4
                    chart.width = 20
                    chart.height = 12
                    ws.add_chart(chart, f"A{fila_res + 2}")
                except:
                    pass

                nombre_arch = f"RAMS_{titulo.replace(' ', '_')[:25]}_{int(time.time())}.xlsx"
                ruta_excel = os.path.join(CARPETA_TRABAJO, nombre_arch)
                wb.save(ruta_excel)

                self.log_to_hud("📊 EXCEL", f"✅ Archivo Excel creado: {nombre_arch}")
                self.hablar(f"Hoja de cálculo lista, Comandante.")
                notification.notify(title="R.A.M.S. Excel", message=f"Excel '{titulo}' creado.", app_name="MAGI")
                os.startfile(ruta_excel)

            except ImportError:
                self.log_to_hud("❌ EXCEL", "Falta openpyxl. Ejecuta: pip install openpyxl")
            except Exception as e:
                self.log_to_hud("❌ EXCEL", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    # ==========================================
    # 📑 CREAR PDF VISUAL
    # ==========================================
    def crear_documento_pdf(self, datos):
        """Crea un PDF visualmente llamativo. datos = 'titulo|contenido'"""
        def _proceso():
            try:
                self.actualizar_estado_hud("CREANDO PDF...", "#8e44ad")
                partes = datos.split("|", 1)
                titulo = partes[0].strip() if partes else "Reporte RAMS"
                contenido_usuario = partes[1].strip() if len(partes) > 1 else datos

                self.log_to_hud("📑 PDF", f"Generando contenido para: {titulo}")

                prompt_pdf = (
                    f"Genera contenido para un PDF profesional sobre: '{contenido_usuario}'.\n"
                    f"Responde SOLO con JSON sin markdown:\n"
                    f'{{"titulo": "...", "subtitulo": "...", "secciones": [{{"titulo": "...", "puntos": ["...","..."]}}]}}\n'
                    f"Al menos 3 secciones con 3-5 puntos cada una."
                )
                try:
                    resp_json = genai.GenerativeModel('gemini-2.5-flash').generate_content(prompt_pdf).text.strip()
                    resp_json = resp_json.replace("```json", "").replace("```", "").strip()
                    estructura = json.loads(resp_json)
                except:
                    estructura = {"titulo": titulo, "subtitulo": "", "secciones": [{"titulo": "Contenido", "puntos": [contenido_usuario]}]}

                from reportlab.lib.pagesizes import A4
                from reportlab.lib import colors
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
                from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

                nombre_arch = f"RAMS_{titulo.replace(' ', '_')[:25]}_{int(time.time())}.pdf"
                ruta_pdf = os.path.join(CARPETA_TRABAJO, nombre_arch)

                doc = SimpleDocTemplate(ruta_pdf, pagesize=A4,
                                        rightMargin=2*cm, leftMargin=2*cm,
                                        topMargin=2*cm, bottomMargin=2*cm)

                # Colores
                ROJO    = colors.HexColor("#FF3B14")
                NARANJA = colors.HexColor("#FF9D00")
                NEGRO   = colors.HexColor("#0A0505")
                VERDE   = colors.HexColor("#39FF14")
                GRIS    = colors.HexColor("#F0F0F0")

                estilos = getSampleStyleSheet()

                estilo_titulo = ParagraphStyle("titulo", fontSize=26, fontName="Helvetica-Bold",
                                               textColor=ROJO, alignment=TA_CENTER, spaceAfter=6)
                estilo_subtit = ParagraphStyle("subtit", fontSize=13, fontName="Helvetica-Oblique",
                                               textColor=NARANJA, alignment=TA_CENTER, spaceAfter=4)
                estilo_fecha  = ParagraphStyle("fecha", fontSize=9, fontName="Helvetica",
                                               textColor=colors.HexColor("#39FF14"), alignment=TA_CENTER, spaceAfter=16)
                estilo_head   = ParagraphStyle("head", fontSize=14, fontName="Helvetica-Bold",
                                               textColor=colors.HexColor("#C03921"), spaceBefore=14, spaceAfter=6)
                estilo_punto  = ParagraphStyle("punto", fontSize=11, fontName="Helvetica",
                                               textColor=NEGRO, spaceBefore=4, spaceAfter=4,
                                               leftIndent=20, alignment=TA_JUSTIFY)
                estilo_pie    = ParagraphStyle("pie", fontSize=9, fontName="Helvetica-Bold",
                                               textColor=ROJO, alignment=TA_CENTER)

                elementos = []

                # Cabecera
                elementos.append(Paragraph(estructura.get("titulo", titulo).upper(), estilo_titulo))
                elementos.append(HRFlowable(width="100%", thickness=2, color=ROJO, spaceAfter=6))
                if estructura.get("subtitulo"):
                    elementos.append(Paragraph(estructura["subtitulo"], estilo_subtit))
                elementos.append(Paragraph(
                    f"Generado por {self.nombre_ia} — R.A.M.S. V2.0 | {time.strftime('%d/%m/%Y %H:%M')}",
                    estilo_fecha
                ))
                elementos.append(Spacer(1, 0.5*cm))

                # Secciones
                for i, seccion in enumerate(estructura.get("secciones", [])):
                    ico = "▶" if i % 2 == 0 else "◆"
                    elementos.append(Paragraph(f"{ico} {seccion.get('titulo','Sección').upper()}", estilo_head))
                    elementos.append(HRFlowable(width="80%", thickness=1, color=colors.HexColor("#C03921"), spaceAfter=4))
                    for punto in seccion.get("puntos", []):
                        elementos.append(Paragraph(f"• {punto}", estilo_punto))
                    elementos.append(Spacer(1, 0.3*cm))

                # Pie
                elementos.append(Spacer(1, 0.5*cm))
                elementos.append(HRFlowable(width="100%", thickness=1, color=ROJO, spaceAfter=6))
                elementos.append(Paragraph(f"━━━ R.A.M.S. V2.0 | {self.nombre_ia} | Confidencial ━━━", estilo_pie))

                doc.build(elementos)

                self.log_to_hud("📑 PDF", f"✅ PDF creado: {nombre_arch}")
                self.hablar(f"El PDF ha sido generado, Comandante.")
                notification.notify(title="R.A.M.S. PDF", message=f"PDF '{titulo}' creado.", app_name="MAGI")
                os.startfile(ruta_pdf)

            except ImportError:
                self.log_to_hud("❌ PDF", "Falta reportlab. Ejecuta: pip install reportlab")
            except Exception as e:
                self.log_to_hud("❌ PDF", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    # ==========================================
    # 🎥 GRABAR CÁMARA
    # ==========================================
    def grabar_camara(self, duracion_seg=15):
        """Graba video desde la cámara web y lo guarda en el Workspace."""
        def _proceso():
            try:
                import cv2
                self.actualizar_estado_hud("GRABANDO CÁMARA...", "#e74c3c")
                self.log_to_hud("🎥 CÁMARA", f"Iniciando grabación de {duracion_seg}s...")
                self.hablar(f"Iniciando grabación de cámara. {duracion_seg} segundos.")

                cap = cv2.VideoCapture(0)
                if not cap.isOpened():
                    self.log_to_hud("❌ CÁMARA", "No se pudo abrir la cámara.")
                    return

                ancho  = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                alto   = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                fps    = int(cap.get(cv2.CAP_PROP_FPS)) or 25

                nombre_arch = f"camara_{int(time.time())}.avi"
                ruta_video  = os.path.join(CARPETA_TRABAJO, nombre_arch)
                fourcc = cv2.VideoWriter_fourcc(*"XVID")
                out = cv2.VideoWriter(ruta_video, fourcc, fps, (ancho, alto))

                inicio = time.time()
                frames_grabados = 0
                while time.time() - inicio < duracion_seg:
                    ret, frame = cap.read()
                    if not ret:
                        break
                    out.write(frame)
                    frames_grabados += 1

                cap.release()
                out.release()

                self.log_to_hud("🎥 CÁMARA", f"✅ Video guardado: {nombre_arch} ({frames_grabados} frames)")
                self.hablar("Grabación de cámara completada.")
                notification.notify(title="R.A.M.S. Cámara", message=f"Video guardado: {nombre_arch}", app_name="MAGI")

            except ImportError:
                self.log_to_hud("❌ CÁMARA", "Falta opencv-python. Ejecuta: pip install opencv-python")
            except Exception as e:
                self.log_to_hud("❌ CÁMARA", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    # ==========================================
    # 🎙️ GRABAR MICRÓFONO (WAV)
    # ==========================================
    def grabar_microfono(self, duracion_seg=10):
        """Graba audio del micrófono y lo guarda en WAV."""
        def _proceso():
            try:
                self.actualizar_estado_hud("GRABANDO MIC...", "#e74c3c")
                self.log_to_hud("🎙️ GRABACIÓN", f"Grabando micrófono: {duracion_seg}s...")
                self.hablar(f"Grabando micrófono por {duracion_seg} segundos.")

                r = sr.Recognizer()
                with sr.Microphone(sample_rate=44100) as source:
                    r.adjust_for_ambient_noise(source, duration=0.5)
                    audio = r.record(source, duration=duracion_seg)

                import wave
                nombre_arch = f"mic_{int(time.time())}.wav"
                ruta_wav = os.path.join(CARPETA_TRABAJO, nombre_arch)
                with wave.open(ruta_wav, "wb") as wf:
                    wf.setnchannels(1)
                    wf.setsampwidth(2)
                    wf.setframerate(44100)
                    wf.writeframes(audio.get_wav_data())

                self.log_to_hud("🎙️ GRABACIÓN", f"✅ Audio guardado: {nombre_arch}")
                self.hablar("Grabación de micrófono completada.")
                notification.notify(title="R.A.M.S. Micrófono", message=f"Audio: {nombre_arch}", app_name="MAGI")

            except Exception as e:
                self.log_to_hud("❌ GRABACIÓN MIC", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    # ==========================================
    # 🎮 BUSCAR Y COMPRAR JUEGOS (Steam + Eneba)
    # ==========================================
    def buscar_juego_steam(self, nombre_juego):
        """Busca el juego en Steam y en Eneba buscando la mejor oferta."""
        def _proceso():
            try:
                self.actualizar_estado_hud("BUSCANDO JUEGO...", "#9b59b6")
                self.log_to_hud("🎮 STEAM", f"Buscando: {nombre_juego}")
                self.hablar(f"Buscando {nombre_juego} en Steam y tiendas de ofertas.")

                nombre_enc = nombre_juego.replace(" ", "+")
                nombre_url = nombre_juego.replace(" ", "%20")

                # Buscar en Steam Store API (gratuita)
                try:
                    steam_url = f"https://store.steampowered.com/api/storesearch/?term={nombre_enc}&l=spanish&cc=CO"
                    resp = requests.get(steam_url, timeout=10)
                    datos = resp.json()
                    items = datos.get("items", [])
                    if items:
                        juego = items[0]
                        precio_steam = juego.get("price", {})
                        nombre_real = juego["name"]
                        app_id = juego["id"]
                        precio_final = precio_steam.get("final", 0) / 100 if precio_steam else 0
                        descuento    = precio_steam.get("discount_percent", 0) if precio_steam else 0

                        self.log_to_hud("🎮 STEAM", f"Encontrado: {nombre_real}")
                        if precio_final > 0:
                            self.log_to_hud("💰 PRECIO STEAM", f"${precio_final:.2f} USD ({descuento}% OFF)")
                        else:
                            self.log_to_hud("🆓 STEAM", "Juego gratuito o precio no disponible.")

                        # Abrir página del juego en Chrome
                        url_juego = f"https://store.steampowered.com/app/{app_id}"
                        abrir_url_chrome(url_juego)
                        self.log_to_hud("🌐 STEAM", f"Página abierta en Chrome: {url_juego}")
                    else:
                        self.log_to_hud("❌ STEAM", f"No encontré '{nombre_juego}' en Steam.")
                        abrir_url_chrome(f"https://store.steampowered.com/search/?term={nombre_enc}")
                except Exception as e:
                    self.log_to_hud("⚠️ STEAM", f"Error API Steam: {e}")
                    abrir_url_chrome(f"https://store.steampowered.com/search/?term={nombre_enc}")

                # Buscar en Eneba (mejores precios)
                time.sleep(1)
                self.log_to_hud("🏪 ENEBA", "Buscando precio en Eneba...")
                eneba_url = f"https://www.eneba.com/latam/store/games?search={nombre_url}"
                abrir_url_chrome(eneba_url)
                self.log_to_hud("🏪 ENEBA", f"Eneba abierto: {eneba_url}")

                # Buscar también en Fanatical
                time.sleep(0.8)
                fanatical_url = f"https://www.fanatical.com/es/search?search={nombre_enc}"
                abrir_url_chrome(fanatical_url)
                self.log_to_hud("🏪 FANATICAL", f"Fanatical abierto para comparar precios.")

                self.hablar(f"He abierto Steam, Eneba y Fanatical para que compares el precio de {nombre_juego}.")
                self.log_to_hud("✅ COMPRAS", "Revisa las pestañas abiertas para comparar precios y comprar.")
                notification.notify(title="R.A.M.S. Tienda", message=f"Buscando: {nombre_juego}", app_name="MAGI")

            except Exception as e:
                self.log_to_hud("❌ STEAM", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    # ==========================================
    # ⌨️ GRABADOR DE INPUTS (Teclado + Mouse + Control)
    # ==========================================
    _grabando_inputs = False
    _log_inputs = []

    def iniciar_grabacion_inputs(self):
        """Graba todos los inputs: teclado, mouse y gamepad."""
        self._grabando_inputs = True
        self._log_inputs = []

        def _grabar():
            try:
                import pynput
                from pynput import keyboard as kb, mouse as ms

                self.log_to_hud("⌨️ INPUTS", "Grabación de inputs iniciada. Di [PARAR_INPUTS] para detener.")
                self.hablar("Grabación de inputs iniciada.")
                self.actualizar_estado_hud("GRABANDO INPUTS...", "#e74c3c")

                def on_key_press(key):
                    if not self._grabando_inputs:
                        return False
                    try:
                        self._log_inputs.append({"tipo": "tecla", "evento": "press", "valor": str(key.char), "t": time.time()})
                    except:
                        self._log_inputs.append({"tipo": "tecla", "evento": "press", "valor": str(key), "t": time.time()})

                def on_key_release(key):
                    if not self._grabando_inputs:
                        return False

                def on_click(x, y, button, pressed):
                    if not self._grabando_inputs:
                        return False
                    if pressed:
                        self._log_inputs.append({"tipo": "mouse_click", "boton": str(button), "x": x, "y": y, "t": time.time()})

                def on_scroll(x, y, dx, dy):
                    if not self._grabando_inputs:
                        return False
                    self._log_inputs.append({"tipo": "mouse_scroll", "dx": dx, "dy": dy, "t": time.time()})

                kb_listener  = kb.Listener(on_press=on_key_press, on_release=on_key_release)
                ms_listener  = ms.Listener(on_click=on_click, on_scroll=on_scroll)
                kb_listener.start()
                ms_listener.start()

                # Intentar gamepad con pygame
                try:
                    pygame.joystick.init()
                    if pygame.joystick.get_count() > 0:
                        pad = pygame.joystick.Joystick(0)
                        pad.init()
                        self.log_to_hud("🎮 CONTROL", f"Gamepad detectado: {pad.get_name()}")
                        while self._grabando_inputs:
                            pygame.event.pump()
                            for i in range(pad.get_numbuttons()):
                                if pad.get_button(i):
                                    self._log_inputs.append({"tipo": "gamepad_btn", "boton": i, "t": time.time()})
                            time.sleep(0.05)
                    else:
                        while self._grabando_inputs:
                            time.sleep(0.5)
                except:
                    while self._grabando_inputs:
                        time.sleep(0.5)

                kb_listener.stop()
                ms_listener.stop()

                # Guardar log
                nombre_log = f"inputs_{int(time.time())}.json"
                ruta_log = os.path.join(CARPETA_TRABAJO, nombre_log)
                with open(ruta_log, "w", encoding="utf-8") as f:
                    json.dump(self._log_inputs, f, ensure_ascii=False, indent=2)

                self.log_to_hud("⌨️ INPUTS", f"✅ Log guardado: {nombre_log} ({len(self._log_inputs)} eventos)")
                self.hablar(f"Grabación detenida. {len(self._log_inputs)} eventos registrados.")
                self.actualizar_estado_hud("STANDBY", "#39ff14")
                notification.notify(title="R.A.M.S. Inputs", message=f"{len(self._log_inputs)} eventos grabados.", app_name="MAGI")

            except ImportError:
                self.log_to_hud("❌ INPUTS", "Falta pynput. Ejecuta: pip install pynput")
                self._grabando_inputs = False
                self.actualizar_estado_hud("STANDBY", "#39ff14")
            except Exception as e:
                self.log_to_hud("❌ INPUTS", str(e))
                self._grabando_inputs = False
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_grabar, daemon=True).start()

    def parar_grabacion_inputs(self):
        """Detiene la grabación de inputs."""
        self._grabando_inputs = False
        self.log_to_hud("⌨️ INPUTS", "Deteniendo grabación de inputs...")

    # ==========================================
    # 📎 ANALIZAR ARCHIVO ADJUNTO (Video/Imagen)
    # ==========================================
    def analizar_archivo_adjunto(self, ruta_archivo):
        """Analiza un video o imagen con Gemini."""
        def _proceso():
            try:
                self.actualizar_estado_hud("ANALIZANDO ARCHIVO...", "#facc15")
                ruta = ruta_archivo.strip().strip('"').strip("'")

                if not os.path.exists(ruta):
                    self.log_to_hud("❌ ADJUNTO", f"Archivo no encontrado: {ruta}")
                    return

                ext = os.path.splitext(ruta)[1].lower()
                nombre = os.path.basename(ruta)

                self.log_to_hud("📎 ADJUNTO", f"Analizando: {nombre}")
                self.hablar(f"Analizando el archivo {nombre}, Comandante.")

                es_video  = ext in [".mp4", ".avi", ".mkv", ".mov", ".webm", ".m4v"]
                es_imagen = ext in [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"]

                if es_imagen:
                    img_subida = genai.upload_file(ruta)
                    prompt = "Analiza esta imagen en detalle. Describe qué ves, elementos importantes, colores, composición y cualquier texto visible. Sé específico y útil."
                    respuesta = self.sesion_chat.send_message([prompt, img_subida]).text.strip()

                elif es_video:
                    self.log_to_hud("📎 ADJUNTO", "Subiendo video a Gemini (puede tardar)...")
                    video_subido = genai.upload_file(ruta)
                    # Esperar a que procese
                    while video_subido.state.name == "PROCESSING":
                        time.sleep(3)
                        video_subido = genai.get_file(video_subido.name)
                    if video_subido.state.name == "FAILED":
                        self.log_to_hud("❌ ADJUNTO", "El video falló al procesarse.")
                        return
                    prompt = "Analiza este video en detalle. Describe qué ocurre, los momentos clave, personas, objetos, acciones y cualquier texto visible. Sé específico."
                    respuesta = self.sesion_chat.send_message([prompt, video_subido]).text.strip()

                else:
                    self.log_to_hud("❌ ADJUNTO", f"Formato no soportado: {ext}")
                    return

                texto_limpio = re.sub(r'```.*?```', '', respuesta, flags=re.DOTALL)
                texto_limpio = re.sub(r'\[.*?\]', '', texto_limpio).replace('*', '').strip()
                if texto_limpio:
                    self.log_to_hud(f"🤖 {self.nombre_ia.upper()} (ADJUNTO)", texto_limpio)
                    self.hablar(texto_limpio[:400])  # Limitar voz

                notification.notify(title="R.A.M.S. Adjunto", message=f"Análisis de {nombre} completado.", app_name="MAGI")

            except Exception as e:
                self.log_to_hud("❌ ADJUNTO", str(e))
            finally:
                self.actualizar_estado_hud("STANDBY", "#39ff14")

        threading.Thread(target=_proceso, daemon=True).start()

    def adjuntar_archivo_dialog(self):
        """Abre un diálogo para seleccionar un archivo y analizarlo."""
        def _abrir():
            import tkinter as tk
            from tkinter import filedialog
            root = tk.Tk()
            root.withdraw()
            ruta = filedialog.askopenfilename(
                title="Selecciona imagen o video para analizar",
                filetypes=[
                    ("Imágenes y Videos", "*.jpg *.jpeg *.png *.gif *.bmp *.webp *.mp4 *.avi *.mkv *.mov *.webm"),
                    ("Imágenes", "*.jpg *.jpeg *.png *.gif *.bmp *.webp"),
                    ("Videos", "*.mp4 *.avi *.mkv *.mov *.webm"),
                    ("Todos los archivos", "*.*"),
                ]
            )
            root.destroy()
            if ruta:
                self.analizar_archivo_adjunto(ruta)
            else:
                self.log_to_hud("📎 ADJUNTO", "Selección cancelada.")
        threading.Thread(target=_abrir, daemon=True).start()


    # ==========================================
    def _obtener_datos_gpu(self):
        """Obtiene uso y temperatura de GPU via GPUtil o WMI."""
        gpu_uso = 0
        gpu_temp = 0
        gpu_nombre = "N/A"
        try:
            import GPUtil
            gpus = GPUtil.getGPUs()
            if gpus:
                gpu = gpus[0]
                gpu_uso = gpu.load * 100
                gpu_temp = gpu.temperature
                gpu_nombre = gpu.name
        except:
            pass
        if gpu_temp == 0:
            try:
                import wmi
                w = wmi.WMI(namespace="root\\OpenHardwareMonitor")
                for sensor in w.Sensor():
                    if sensor.SensorType == "Temperature" and "GPU" in sensor.Name:
                        gpu_temp = sensor.Value
                        break
            except:
                pass
        return gpu_uso, gpu_temp, gpu_nombre

    def _obtener_temp_cpu(self):
        """Obtiene temperatura de CPU."""
        try:
            import wmi
            w = wmi.WMI(namespace="root\\OpenHardwareMonitor")
            for sensor in w.Sensor():
                if sensor.SensorType == "Temperature" and "CPU" in sensor.Name:
                    return sensor.Value
        except:
            pass
        try:
            temps = psutil.sensors_temperatures()
            if temps:
                for name, entries in temps.items():
                    if entries:
                        return entries[0].current
        except:
            pass
        return 0

    def telemetria_loop(self):
        while NUCLEO_VIVO:
            if self.window:
                try:
                    cpu = psutil.cpu_percent()
                    ram = psutil.virtual_memory().percent
                    gpu_uso, gpu_temp, gpu_nombre = self._obtener_datos_gpu()
                    cpu_temp = self._obtener_temp_cpu()
                    self.window.evaluate_js(
                        f"updateTelemetry({cpu}, {ram}, {gpu_uso:.1f}, {gpu_temp:.1f}, {cpu_temp:.1f}, '{gpu_nombre}');"
                    )
                except:
                    pass
            time.sleep(2)

    # ==========================================
    # 👂 ESCUCHA CONTINUA POR PALABRA CLAVE
    # ==========================================
    def oido_siempre_alerta(self):
        r = sr.Recognizer()
        with sr.Microphone() as source:
            while NUCLEO_VIVO:
                if not self.oido_activo:
                    time.sleep(1)
                    continue
                try:
                    audio = r.listen(source, timeout=None, phrase_time_limit=8)
                    texto = r.recognize_google(audio, language="es-ES").lower()
                    if self.nombre_ia.lower() in texto:
                        comando = texto.replace(self.nombre_ia.lower(), "").strip()
                        self.log_to_hud("🎤 COMANDO DE VOZ", texto)
                        if comando:
                            threading.Thread(target=self.ejecutar_logica, args=(comando,), daemon=True).start()
                except:
                    continue


# ==========================================
# 🚀 PROTOCOLO DE ARRANQUE
# ==========================================
def iniciar_hilos_background(window, api_bridge):
    api_bridge.window = window
    time.sleep(1.5)
    api_bridge.log_to_hud("⚙️ SISTEMA", "Enlace Sincronizado. MAGI Terminal Operativo.")
    threading.Thread(target=api_bridge.telemetria_loop, daemon=True).start()
    threading.Thread(target=api_bridge.oido_siempre_alerta, daemon=True).start()
    threading.Thread(target=api_bridge.recordatorio_loop, daemon=True).start()
    threading.Thread(target=api_bridge.capsula_loop, daemon=True).start()
    # Verificar cápsulas/recordatorios pendientes al arrancar
    api_bridge.listar_recordatorios()


if __name__ == "__main__":
    if not os.path.exists(ARCHIVO_CONFIG):
        app_config = VentanaConfiguracion()
        app_config.mainloop()

    try:
        with open(ARCHIVO_CONFIG, "r") as f:
            lineas = f.read().splitlines()
            nombre_ia   = lineas[0]
            api_key     = lineas[1]
            mi_correo   = lineas[2] if len(lineas) > 2 else None
            mi_pass_app = lineas[3] if len(lineas) > 3 else None
    except Exception as e:
        print(f"Error fatal al leer configuración: {e}")
        exit()

    puente_magi = RamsBridgeAPI(nombre_ia, api_key, mi_correo, mi_pass_app)
    ruta_html = os.path.abspath(ARCHIVO_HUD)

    ventana_magi = webview.create_window(
        title=f"MAGI TERMINAL - {nombre_ia}",
        url=f"file://{ruta_html}",
        js_api=puente_magi,
        frameless=False,
        easy_drag=False,
        width=1250, height=780,
        background_color='#0a0505'
    )
    webview.start(iniciar_hilos_background, args=(ventana_magi, puente_magi), gui='edgechromium', debug=False)

    # Al cerrar la ventana, escribir la entrada del diario
    puente_magi.escribir_entrada_diario()